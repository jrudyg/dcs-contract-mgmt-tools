#!/usr/bin/env python3
"""
anonymize.py — P3 contract anonymization pipeline (detect → review → apply)

Layers applied in order:
  1. Counterparty redaction  (deterministic map from mapping.json)
  2. PII detection/redaction (Presidio + spaCy en_core_web_lg)
  3. Commercial terms regex  (currency, percentages, payment terms, rates)

Two-phase model:
  detect_file()  — extract text, find spans across all three layers, write a
                   <stem>.review.json with confirmed=null for human review.
                   Does NOT write .anon.txt.
  apply_file()   — read a reviewed <stem>.review.json, replay the three layers
                   applying ONLY confirmed=true spans, write <stem>.anon.txt and
                   <stem>.audit.json.

Output per file (written to --output-dir, gitignored):
  <stem>.review.json — detected spans with confirmed flag (detect phase)
  <stem>.anon.txt    — redacted plain text                (apply phase)
  <stem>.audit.json  — redaction inventory (confirmed spans only, apply phase)

Usage:
  python anonymize.py --input file.docx --detect-only
  python anonymize.py --input file.docx --full-run --audit
  python anonymize.py --input path/to/dir/ --output-dir anonymization/output --audit
  python anonymize.py --input file.pdf --rebuild-map
"""

import argparse
import json
import logging
import pathlib
import re
import time
from typing import Optional

import pdfplumber
from rapidfuzz import fuzz
from docx import Document
from presidio_analyzer import AnalyzerEngine
from presidio_analyzer import RecognizerResult
from presidio_anonymizer import AnonymizerEngine
from presidio_anonymizer.entities import OperatorConfig

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

EXCLUDED_COUNTERPARTIES: frozenset[str] = frozenset({
    "(blank)",
    "A. MNDA Template",
    "ACI for Signature",
})

PRESIDIO_ENTITIES: list[str] = [
    "PERSON",
    "EMAIL_ADDRESS",
    "PHONE_NUMBER",
    "US_SSN",
    "DATE_TIME",
    "LOCATION",
    "NRP",
    "US_BANK_NUMBER",
    "CREDIT_CARD",
]

# Commercial terms patterns — applied in listed order, IGNORECASE|MULTILINE.
COMMERCIAL_PATTERNS: list[tuple[re.Pattern, str]] = [
    (re.compile(r'\$[\d,]+(?:\.\d{1,2})?',  re.MULTILINE), '[AMOUNT]'),
    (re.compile(r'\d+(?:\.\d+)?\s*%',        re.MULTILINE), '[PERCENT]'),
    (re.compile(r'\bnet\s*\d+\b',            re.IGNORECASE | re.MULTILINE), '[PAYMENT_TERM]'),
    (re.compile(r'\$[\d.]+\s*/\s*\w+',       re.MULTILINE), '[RATE]'),
    (re.compile(
        r'\b\d[\d,]*\.\d{2}\s*(?:USD|EUR|GBP|CAD)\b',
        re.IGNORECASE | re.MULTILINE,
    ), '[AMOUNT]'),
    (re.compile(
        r'\b(?:USD|EUR|GBP|CAD)\s*\d[\d,]*\.\d{2}\b',
        re.IGNORECASE | re.MULTILINE,
    ), '[AMOUNT]'),
]

# B2.1 — TERM_DURATION context-gated reintroduction
# Fires only when a duration expression sits within 150 chars of a
# payment/fee/billing keyword. Warranty/delivery durations excluded.
TERM_DURATION_PATTERN: re.Pattern = re.compile(
    r'\b\d+\s*(?:days?|weeks?|months?|years?)\b',
    re.IGNORECASE,
)

TERM_DURATION_CONTEXT: re.Pattern = re.compile(
    r'\b(?:payment|fee|fees|invoice|invoicing|billing|net|due|'
    r'remit|remittance|payable|subscription|renewal|term)\b',
    re.IGNORECASE,
)

TERM_DURATION_WINDOW: int = 150  # chars each side of match to search

# B0.5 — Commercial context tagging
# Tags each commercial span with a context_type based on surrounding text.
# Does not change detection — annotation only for review efficiency.
COMMERCIAL_CONTEXT_PATTERNS: list[tuple[re.Pattern, str]] = [
    (re.compile(
        r'\b(?:payment|invoice|due|net|billing|remit|payable)\b',
        re.IGNORECASE), 'payment'),
    (re.compile(
        r'\b(?:liability|liabilities|indemnif|cap|exceed|limit|ceiling)\b',
        re.IGNORECASE), 'liability_cap'),
    (re.compile(
        r'\b(?:insurance|insur|coverage|policy|policies|premium)\b',
        re.IGNORECASE), 'insurance'),
]

COMMERCIAL_CONTEXT_WINDOW: int = 200  # chars to search on each side of span

# B0.6 — Contract-type detection tuning
# Contract type seeds commercial scanning weight only.
# NDA → skip commercial layer entirely (NDAs rarely contain pricing).
# SOW / PO / MSA / unknown → full commercial scanning (default).
CONTRACT_TYPE_SKIP_COMMERCIAL: frozenset[str] = frozenset({
    'nda', 'mutual nda', 'mnda', 'non-disclosure', 'non disclosure',
})

SUPPORTED_EXTENSIONS: frozenset[str] = frozenset({'.pdf', '.docx', '.doc', '.txt'})

# Alias-pass stopwords — generic role/duration/filler words that must never be
# treated as a counterparty alias even if they appear in alias_map.
ALIAS_STOPWORDS: frozenset[str] = frozenset({
    # Generic legal/commercial role labels — not party identities
    'buyer', 'seller', 'vendor', 'supplier', 'contractor', 'subcontractor',
    'client', 'customer', 'company', 'party', 'parties', 'owner',
    'consultant', 'agent', 'distributor', 'reseller', 'licensor', 'licensee',
    # Duration / quantity words that slip through alias detection
    'days', 'weeks', 'months', 'years', 'term',
    # Common false-positive single words
    'the', 'and', 'or', 'for', 'not', 'any', 'all',
    # Common nouns that collide with short single-word aliases (E19)
    'capacity', 'api',
})

# Generic words / section headings / role labels Presidio frequently mis-tags as
# PERSON/LOCATION/NRP/DATE_TIME. Suppress these presidio spans (tuning 2026-06-20).
PRESIDIO_STOPWORDS: frozenset[str] = ALIAS_STOPWORDS | frozenset({
    'deliverables', 'monthly', 'the later', 'project overview', 'overview',
    'timeline', 'personnel', 'exclusions', 'scope of services', 'payment terms',
    'miscellaneous', 'governing law', 'fees',
})

# Common single dictionary words that must never count as counterparty identity
# in the D1/D2 leakage/sweep scans. Keep in sync with build_map.py COMMON_WORDS.
COMMON_WORDS: frozenset[str] = frozenset({
    "material", "general", "process", "system", "systems", "service", "services",
    "group", "global", "national", "american", "united", "standard", "industrial",
    "industries", "solutions", "technology", "technologies", "network", "networks",
    "management", "resources", "advanced", "integrated", "international", "design",
    "designed", "conveyor", "conveyors", "distribution", "logistics", "supply",
    "capital", "partners", "holdings", "enterprises", "corp", "corporation",
    "company", "limited", "inc", "llc", "ltd", "co", "the", "of", "and",
    "tractor", "action", "serve", "merc",
})

# B0.3 — Structure detection constants
# Lines where >80% of alphabetic chars are uppercase → structural header.
# Lines matching numbered-section pattern → structural.
# Signature labels at line start → structural.
STRUCTURAL_UPPERCASE_THRESHOLD: float = 0.80

STRUCTURAL_NUMBERED = re.compile(
    r'^\s*(\d+\.)+\s',
    re.MULTILINE,
)

SIGNATURE_LABELS: frozenset[str] = frozenset({
    'by:', 'name:', 'title:', 'date:', 'signature:',
})

# B0.4 — Signature-block zone detection
# PII suppression inside zones: LOCATION, NRP, DATE_TIME suppressed;
# PERSON retained (names on signature lines must be redacted).
SIGNATURE_ZONE_LABELS: frozenset[str] = frozenset({
    'by:', 'name:', 'title:', 'date:',
})

SIGNATURE_ZONE_SUPPRESS: frozenset[str] = frozenset({
    'LOCATION', 'NRP', 'DATE_TIME',
})

SIGNATURE_ZONE_WINDOW: int = 10  # lines within which labels must co-occur

# ---------------------------------------------------------------------------
# Logging
# ---------------------------------------------------------------------------

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s  %(levelname)-7s  %(message)s',
    datefmt='%H:%M:%S',
)
log = logging.getLogger(__name__)

# Robust logging on a cp1252 host. The scheduled-task server runs under pythonw
# whose stderr is cp1252, which cannot encode characters like '→' or smart quotes
# that appear in log messages/data. Without this, a single such char raises
# UnicodeEncodeError out of the logging call (and again inside handleError) and
# surfaces as an HTTP 500. backslashreplace keeps the stream writable; disabling
# raiseExceptions ensures any residual emit error can never abort a request.
for _h in logging.getLogger().handlers:
    _stream = getattr(_h, 'stream', None)
    if _stream is not None and hasattr(_stream, 'reconfigure'):
        try:
            _stream.reconfigure(errors='backslashreplace')
        except Exception:
            pass
logging.raiseExceptions = False

# ---------------------------------------------------------------------------
# Mapping: build or load
# ---------------------------------------------------------------------------

def build_mapping(
    counterparties_path: pathlib.Path,
    mapping_path: pathlib.Path,
) -> tuple[dict[str, str], dict[str, str]]:
    """Parse kb/COUNTERPARTIES.md and write a v1 mapping.json (no aliases).

    For a full v2 map with alias detection run build_map.py directly.
    Returns (name_map, alias_map); alias_map is always empty from this builder.
    """
    name_pattern = re.compile(r'^\*\*(.+?)\*\*')
    entries: dict[str, str] = {}
    idx = 1
    with open(counterparties_path, encoding='utf-8') as fh:
        for line in fh:
            m = name_pattern.match(line)
            if m:
                name = m.group(1).strip()
                if name not in EXCLUDED_COUNTERPARTIES:
                    entries[name] = f'PARTY-{idx:04d}'
                    idx += 1

    payload = {
        'generated': time.strftime('%Y-%m-%dT%H:%M:%S'),
        'version': 1,
        'excluded': sorted(EXCLUDED_COUNTERPARTIES),
        'entries': entries,
    }
    mapping_path.parent.mkdir(parents=True, exist_ok=True)
    with open(mapping_path, 'w', encoding='utf-8') as fh:
        json.dump(payload, fh, indent=2, ensure_ascii=False)
    log.info('Built mapping v1: %d entries → %s', len(entries), mapping_path)
    return entries, {}


def load_mapping(
    mapping_path: pathlib.Path,
) -> tuple[dict[str, str], dict[str, str]]:
    """Load mapping.json (v1 or v2) → (name_map, alias_map).

    v1: entries is {name: pseudonym}  — alias_map returned empty.
    v2: entries is {name: {pseudonym, aliases, ...}} — alias_map populated.
    """
    with open(mapping_path, encoding='utf-8') as fh:
        payload = json.load(fh)

    version = payload.get('version', 1)
    raw_entries = payload['entries']

    if version >= 2:
        name_map: dict[str, str] = {
            name: data['pseudonym']
            for name, data in raw_entries.items()
        }
        alias_map: dict[str, str] = {}
        for name, data in raw_entries.items():
            pseudo = data['pseudonym']
            for alias in data.get('aliases', []):
                if len(alias) >= 3:
                    alias_map[alias] = pseudo
    else:
        name_map = dict(raw_entries)
        alias_map = {}

    log.info(
        'Loaded mapping v%d: %d names, %d aliases from %s',
        version, len(name_map), len(alias_map), mapping_path,
    )
    return name_map, alias_map


def get_mapping(
    mapping_path: pathlib.Path,
    counterparties_path: pathlib.Path,
    rebuild: bool,
) -> tuple[dict[str, str], dict[str, str]]:
    """Return (name_map, alias_map), building mapping.json if necessary."""
    if rebuild:
        if not counterparties_path.exists():
            raise FileNotFoundError(
                f'Cannot build mapping: {counterparties_path} not found.'
            )
        return build_mapping(counterparties_path, mapping_path)
    if not mapping_path.exists():
        if mapping_path.is_absolute():
            raise FileNotFoundError(
                f"Mapping file not found at explicit path: {mapping_path}\n"
                f"Run with --rebuild-map to regenerate, or restore the file from backup."
            )
        if not counterparties_path.exists():
            raise FileNotFoundError(
                f'Cannot build mapping: {counterparties_path} not found.'
            )
        return build_mapping(counterparties_path, mapping_path)
    return load_mapping(mapping_path)

# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------

def extract_text(path: pathlib.Path) -> str:
    """Extract plain text from a PDF, DOCX/DOC, or TXT file."""
    suffix = path.suffix.lower()
    if suffix == '.pdf':
        return _extract_pdf(path)
    if suffix in ('.docx', '.doc'):
        return _extract_docx(path)
    if suffix == '.txt':
        return _extract_txt(path)
    raise ValueError(f'Unsupported file type: {suffix}')


def _extract_pdf(path: pathlib.Path) -> str:
    pages: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
    return '\n'.join(pages)


def _extract_docx(path: pathlib.Path) -> str:
    """Extract text from a .docx, including regions python-docx's
    ``doc.paragraphs`` silently omits: tables, headers/footers, and text
    boxes/shapes. Any counterparty name living only in those regions would
    otherwise be invisible to detection — a confirmed redaction-leakage
    vector (e.g. 'Colmac' x4 sat in a body text box and extracted as 0).
    """
    from docx.oxml.ns import qn

    doc = Document(str(path))
    parts: list[str] = []

    def _table_text(tbl) -> list[str]:
        out: list[str] = []
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text:
                        out.append(p.text)
                for nested in cell.tables:
                    out.extend(_table_text(nested))
        return out

    # Body paragraphs + tables
    for para in doc.paragraphs:
        if para.text:
            parts.append(para.text)
    for tbl in doc.tables:
        parts.extend(_table_text(tbl))

    # Headers / footers (primary, first-page, even-page) for each section
    for section in doc.sections:
        for hf in (
            section.header, section.footer,
            section.first_page_header, section.first_page_footer,
            section.even_page_header, section.even_page_footer,
        ):
            if hf is None:
                continue
            for para in hf.paragraphs:
                if para.text:
                    parts.append(para.text)
            for tbl in hf.tables:
                parts.extend(_table_text(tbl))

    # Text boxes / shapes — w:txbxContent is skipped by paragraph iteration.
    # Walk body + header/footer parts. Dedupe identical blocks to avoid the
    # mc:AlternateContent Choice/Fallback double-count.
    seen_txbx: set[str] = set()
    roots = [doc.element]
    for section in doc.sections:
        for hf in (section.header, section.footer):
            if hf is not None:
                roots.append(hf._element)
    for root in roots:
        for txbx in root.iter(qn('w:txbxContent')):
            text = ''.join(t.text for t in txbx.iter(qn('w:t')) if t.text)
            if text and text not in seen_txbx:
                seen_txbx.add(text)
                parts.append(text)

    return '\n'.join(parts)


def _extract_txt(path: pathlib.Path) -> str:
    for enc in ('utf-8-sig', 'utf-8', 'latin-1'):
        try:
            return path.read_text(encoding=enc)
        except UnicodeDecodeError:
            continue
    raise ValueError(f'Could not decode {path} with utf-8-sig, utf-8, or latin-1')

# ---------------------------------------------------------------------------
# Layer 2: Presidio PII
# ---------------------------------------------------------------------------

_analyzer_instance: Optional[AnalyzerEngine] = None
_anonymizer_instance: Optional[AnonymizerEngine] = None


def _presidio_engines() -> tuple[AnalyzerEngine, AnonymizerEngine]:
    global _analyzer_instance, _anonymizer_instance
    if _analyzer_instance is None:
        log.info('Initialising Presidio engines (first call)…')
        _analyzer_instance = AnalyzerEngine()
        _anonymizer_instance = AnonymizerEngine()
    return _analyzer_instance, _anonymizer_instance

# ---------------------------------------------------------------------------
# Detect phase
# ---------------------------------------------------------------------------

def _score_tier(layer: str, entity_type: str, score: float) -> str:
    """Classify a span into a review priority tier.

    Only presidio spans can fall below 'standard'; all other layers
    (counterparty, counterparty_alias, commercial_regex) are always
    'standard' regardless of their score.
    """
    if layer == 'presidio':
        if score < 0.35:
            return 'low'
        if score < 0.60 and entity_type == 'PERSON':
            return 'medium'
    return 'standard'


BINDING_PATTERN = re.compile(
    r'\(\s*(?:the\s+|hereinafter\s+)?'
    r'["“]([A-Z][A-Za-z0-9 &]{1,40})["”]\s*\)',
    re.IGNORECASE,
)

ROLE_LABELS: frozenset[str] = frozenset({
    'buyer', 'seller', 'company', 'supplier', 'contractor',
    'subcontractor', 'client', 'customer', 'vendor', 'owner',
    'consultant', 'agent', 'distributor', 'reseller',
    'licensor', 'licensee', 'purchaser',
})


def extract_defined_term_bindings(
    text: str,
    name_map: dict[str, str],
) -> dict[str, str]:
    """Parse the first 1,000 characters for defined-term bindings.

    Matches patterns like:
      ("DCS")  (the "Company")  (hereinafter "Supplier")  (“Buyer”)

    For each binding:
      - Finds the nearest preceding capitalized name (within 60 chars) that
        exists in name_map.
      - If found: adds the abbreviation as a document-local alias mapped to
        that party's pseudonym token. Document-local aliases override the
        global alias map and are immune to ALIAS_STOPWORDS.
      - Role-label binding terms (BUYER, SELLER, COMPANY…) are skipped — they
        are structural, not party identities, and must never become aliases.

    Returns:
      local_aliases : dict[str, str]  alias -> pseudonym (to merge with alias_map)
    """
    local_aliases: dict[str, str] = {}

    window = text[:1000]
    for match in BINDING_PATTERN.finditer(window):
        binding_term = match.group(1).strip()
        if binding_term.lower() in ROLE_LABELS:
            continue
        # Nearest preceding known party name within the 60 chars before the
        # binding (rfind -> largest index wins = closest to the binding).
        preceding = window[max(0, match.start() - 60):match.start()].lower()
        best_pos = -1
        best_token: Optional[str] = None
        for name, token in name_map.items():
            pos = preceding.rfind(name.lower())
            if pos > best_pos:
                best_pos = pos
                best_token = token
        if best_token is not None:
            local_aliases[binding_term] = best_token

    return local_aliases


def _compute_signature_zones(text: str) -> list[tuple[int, int]]:
    """Return (start, end) char ranges that are signature-block zones.

    A signature zone is any region where at least 2 of the SIGNATURE_ZONE_LABELS
    ('by:', 'name:', 'title:', 'date:') appear within SIGNATURE_ZONE_WINDOW
    lines of each other.

    Within these zones, LOCATION / NRP / DATE_TIME presidio detections are
    suppressed — they are boilerplate (state of incorporation, signing date)
    not sensitive identity. PERSON detections are retained.

    Returns ranges sorted by start position, merged where overlapping.
    """
    lines = text.splitlines(keepends=True)
    # Build list of (line_index, char_start, char_end, label) for each label hit
    label_hits: list[tuple[int, int, int, str]] = []
    pos = 0
    for i, line in enumerate(lines):
        stripped = line.strip().lower()
        for label in SIGNATURE_ZONE_LABELS:
            if stripped.startswith(label):
                label_hits.append((i, pos, pos + len(line), label))
                break
        pos += len(line)

    if len(label_hits) < 2:
        return []

    # Sliding window: for each hit, check whether ≥1 other hit falls within
    # SIGNATURE_ZONE_WINDOW lines. If so, the span from earliest to latest
    # hit in the window is a zone.
    zones: list[tuple[int, int]] = []
    for i, (li, ls, le, _) in enumerate(label_hits):
        window_hits = [
            h for h in label_hits
            if abs(h[0] - li) <= SIGNATURE_ZONE_WINDOW
        ]
        if len(window_hits) >= 2:
            zone_start = min(h[1] for h in window_hits)
            zone_end   = max(h[2] for h in window_hits)
            zones.append((zone_start, zone_end))

    if not zones:
        return []

    # Merge overlapping zones
    zones.sort()
    merged: list[tuple[int, int]] = [zones[0]]
    for start, end in zones[1:]:
        if start <= merged[-1][1]:
            merged[-1] = (merged[-1][0], max(merged[-1][1], end))
        else:
            merged.append((start, end))
    return merged


def _compute_structural_ranges(text: str) -> list[tuple[int, int]]:
    """Return a list of (start, end) char ranges that are structural.

    Structural = section headers, numbered articles, signature labels.
    These ranges are excluded from span detection UNLESS they already
    contain a counterparty match (B0.3 ordering constraint).

    Rule 1 — uppercase line: >80% of alphabetic chars are uppercase.
    Rule 2 — numbered section: line matches ^\\s*(\\d+\\.)+\\s
    Rule 3 — signature label: line starts with a SIGNATURE_LABELS member
              (case-insensitive, stripped).

    Returns ranges sorted by start position.
    """
    ranges: list[tuple[int, int]] = []
    pos = 0
    for line in text.splitlines(keepends=True):
        line_start = pos
        line_end = pos + len(line)
        stripped = line.strip()

        alpha = [c for c in stripped if c.isalpha()]
        is_uppercase_line = (
            len(alpha) >= 4
            and len(stripped) <= 120
            and sum(1 for c in alpha if c.isupper()) / len(alpha)
            >= STRUCTURAL_UPPERCASE_THRESHOLD
        )

        is_numbered = bool(STRUCTURAL_NUMBERED.match(line))

        is_sig_label = any(
            stripped.lower().startswith(label)
            for label in SIGNATURE_LABELS
        )

        if is_uppercase_line or is_numbered or is_sig_label:
            ranges.append((line_start, line_end))

        pos = line_end

    return ranges


def _commercial_context(
    text: str,
    char_start: int,
    char_end: int,
) -> str:
    """Return a context_type label for a commercial span.

    Searches COMMERCIAL_CONTEXT_WINDOW chars on each side of the span
    for context keywords. First match in COMMERCIAL_CONTEXT_PATTERNS wins.
    Returns 'other' if no pattern matches.
    """
    window_start = max(0, char_start - COMMERCIAL_CONTEXT_WINDOW)
    window_end   = min(len(text), char_end + COMMERCIAL_CONTEXT_WINDOW)
    context      = text[window_start:window_end]

    for pattern, label in COMMERCIAL_CONTEXT_PATTERNS:
        if pattern.search(context):
            return label
    return 'other'


def _sanitize_stem(
    stem: str,
    name_map: dict[str, str],
    alias_map: dict[str, str],
) -> str:
    """Replace counterparty names and aliases in a filename stem with tokens.

    Applies name_map (full names, longest-first) then alias_map (aliases,
    longest-first, word-boundary aware) against the stem string. Same
    replacement logic as Layer 1 detection but against the filename, not
    the document body.

    Falls back to a hash-based stem if the result is empty or contains only
    separators/whitespace after sanitization.

    Returns the sanitized stem (safe to use as an output filename base).
    """
    import hashlib
    result = stem

    # Pass 1: full counterparty names (longest-first, case-insensitive)
    for name, token in sorted(name_map.items(), key=lambda kv: -len(kv[0])):
        result = re.sub(re.escape(name), token, result, flags=re.IGNORECASE)

    # Pass 2: aliases (longest-first, word-boundary, min 3 chars)
    for alias, token in sorted(alias_map.items(), key=lambda kv: -len(kv[0])):
        if len(alias) < 3:
            continue
        result = re.sub(
            r'\b' + re.escape(alias) + r'\b', token, result, flags=re.IGNORECASE
        )

    # Collapse whitespace to underscore; preserve existing hyphens/underscores
    # (so PARTY-XXXX tokens keep their hyphen) and strip leading/trailing junk.
    result = re.sub(r'\s+', '_', result).strip('_')

    # Fallback: if stem is now empty or collision-prone (all tokens, no text)
    if not result or result.replace('_', '').replace('-', '') == '':
        result = 'anon_' + hashlib.sha256(stem.encode()).hexdigest()[:8]

    return result


def _term_duration_spans(text: str) -> list[dict]:
    """Return TERM_DURATION spans gated on nearby payment/billing context.

    A duration match is included only if TERM_DURATION_CONTEXT fires
    within TERM_DURATION_WINDOW chars on either side of the match.
    Warranty, delivery, and other non-commercial durations are excluded
    because they lack the payment keyword neighbours.

    Returns a list of span dicts compatible with the com_matches schema.
    """
    spans: list[dict] = []
    for match in TERM_DURATION_PATTERN.finditer(text):
        window_start = max(0, match.start() - TERM_DURATION_WINDOW)
        window_end   = min(len(text), match.end() + TERM_DURATION_WINDOW)
        context      = text[window_start:window_end]
        if TERM_DURATION_CONTEXT.search(context):
            spans.append({
                'layer':         'commercial_regex',
                'entity_type':   '[TERM_DURATION]',
                'token':         '[TERM_DURATION]',
                'char_start':    match.start(),
                'char_end':      match.end(),
                'original_text': match.group(),
                'score':         1.0,
                'context_type':  'payment',
            })
    return spans


DECISIONS_LIBRARY_PATH: pathlib.Path = pathlib.Path(
    r'C:\Users\jrudy\OneDrive - Diakonia Group, LLC'
    r'\Contract Management - SharePoint\_anon-private\decisions_library.json'
)

# Private-zone layout (outside the Copilot-indexed document tree). Sidecars that
# contain originals or reverse-map data live here, never next to the source.
PRIVATE_DIR: pathlib.Path = DECISIONS_LIBRARY_PATH.parent          # _anon-private\
SHAREPOINT_ROOT: pathlib.Path = PRIVATE_DIR.parent                # SharePoint root
REVIEWS_DIR: pathlib.Path = PRIVATE_DIR / 'reviews'               # *.review.json
VERIFY_DIR: pathlib.Path = PRIVATE_DIR / 'verify'                 # *.verify.json
QUARANTINE_DIR: pathlib.Path = PRIVATE_DIR / 'quarantine'         # failed *.anon.txt
AUDITS_DIR: pathlib.Path = PRIVATE_DIR / 'audits'                 # *.audit.json


def _private_stem(source) -> str:
    """Path-context-unique stem for a source's private sidecars.

    Derived from the source path relative to the SharePoint root, with path
    separators and whitespace collapsed to underscores, so two contracts that
    share a filename in different subfolders never collide
    (e.g. '05-In-Process_Colmac_Colmac-NDA'). Falls back to the bare filename
    stem when the source lies outside the SharePoint root (ad-hoc CLI runs).
    """
    source = pathlib.Path(source)
    try:
        stem_path = str(source.resolve().relative_to(SHAREPOINT_ROOT).with_suffix(''))
    except (ValueError, OSError):
        stem_path = source.stem
    stem = re.sub(r'[\\/\s]+', '_', stem_path).strip('_')
    return stem or source.stem


def review_path_for(source) -> pathlib.Path:
    """Private location of a source's .review.json."""
    return REVIEWS_DIR / f'{_private_stem(source)}.review.json'


def quarantine_path_for(source) -> pathlib.Path:
    """Private quarantine location for a source's failed .anon.txt."""
    return QUARANTINE_DIR / f'{_private_stem(source)}.anon.txt'


def audit_path_for(doc_dir, sanitized_stem: str) -> pathlib.Path:
    """Private location of a shippable output's .audit.json.

    Keyed by the OUTPUT path-context stem — the doc-tree directory (relative to
    the SharePoint root) joined with the sanitized output stem, separators and
    whitespace collapsed to underscores. Same "relative-path → underscores"
    format as reviews/verify, but computed from the shippable .anon.txt path
    rather than the source: de_anonymize is only ever handed the (sanitized)
    .anon.txt, so the audit key must be derivable from it. The audit payload
    still carries `private_stem`, which then locates the source-keyed review.json.
    """
    base = sanitized_stem
    try:
        rel = pathlib.Path(doc_dir).resolve().relative_to(SHAREPOINT_ROOT)
        base = str(rel / sanitized_stem)
    except (ValueError, OSError):
        base = sanitized_stem
    stem = re.sub(r'[\\/\s]+', '_', base).strip('_') or sanitized_stem
    return AUDITS_DIR / f'{stem}.audit.json'


def _load_decisions_library() -> dict:
    """Load decisions_library.json → dict keyed by 'original_text|entity_type'.

    Initialises the file with the canonical schema if it is absent,
    empty, or corrupt. Never raises — returns an empty decisions dict
    on any read failure.

    Schema:
    {
      "version": 1,
      "updated": "<ISO timestamp>",
      "decisions": {
        "OSHA|NRP": {
          "original_text": "OSHA",
          "entity_type": "NRP",
          "decision": "reject",
          "count": 7,
          "last_seen": "2026-06-15"
        }
      }
    }
    Key: f"{original_text}|{entity_type_class}"
    entity_type_class for counterparty/alias spans = "counterparty";
    for presidio spans = the Presidio entity_type string (e.g. "PERSON");
    for commercial spans = the token string (e.g. "[AMOUNT]").
    """
    _EMPTY = {"version": 1, "updated": "", "decisions": {}}
    try:
        if not DECISIONS_LIBRARY_PATH.exists():
            _save_decisions_library(_EMPTY)
            return {}
        raw = DECISIONS_LIBRARY_PATH.read_text(encoding="utf-8").strip()
        if not raw or raw == "{}":
            _save_decisions_library(_EMPTY)
            return {}
        payload = json.loads(raw)
        return payload.get("decisions", {})
    except Exception as exc:
        log.warning("Decisions library unreadable — starting empty: %s", exc)
        return {}


def _save_decisions_library(decisions_or_payload: dict) -> None:
    """Write decisions dict (or full payload) back to disk atomically.

    Accepts either:
      - a bare decisions dict  (keys are "text|type" strings) — wraps it
      - a full payload dict    (has "version" key) — writes as-is

    Uses a .tmp write + rename for atomicity. Never raises — logs on failure.
    """
    try:
        if "version" in decisions_or_payload:
            payload = decisions_or_payload
            payload["updated"] = time.strftime("%Y-%m-%dT%H:%M:%S")
        else:
            payload = {
                "version":   1,
                "updated":   time.strftime("%Y-%m-%dT%H:%M:%S"),
                "decisions": decisions_or_payload,
            }
        tmp = DECISIONS_LIBRARY_PATH.with_suffix(".tmp")
        tmp.write_text(
            json.dumps(payload, indent=2, ensure_ascii=False),
            encoding="utf-8",
        )
        tmp.replace(DECISIONS_LIBRARY_PATH)
    except Exception as exc:
        log.error("Failed to save decisions library: %s", exc)


def append_override_entry(entry: dict) -> None:
    """Append a record to the decisions_library.json `overrides` list.

    Used when a human promotes a quarantined (verification-failed) .anon.txt
    into the shippable document tree — the override is logged for audit.
    Never raises — logs on failure.
    """
    try:
        if DECISIONS_LIBRARY_PATH.exists():
            payload = json.loads(
                DECISIONS_LIBRARY_PATH.read_text(encoding="utf-8").strip() or "{}"
            )
        else:
            payload = {}
        payload.setdefault("version", 1)
        payload.setdefault("decisions", {})
        payload.setdefault("overrides", [])
        payload["overrides"].append(entry)
        _save_decisions_library(payload)  # has "version" → written as full payload
    except Exception as exc:
        log.error("Failed to append override entry: %s", exc)


def _library_key(original_text: str, entity_type: str) -> str:
    """Return the decisions library lookup key for a span."""
    return f"{original_text}|{entity_type}"


def detect_file(
    source: pathlib.Path,
    mapping: dict[str, str],
    alias_map: Optional[dict[str, str]],
    output_dir: pathlib.Path,
    contract_type: str = 'unknown',
    context: Optional[dict] = None,
) -> list[dict]:
    """Detect candidate redaction spans across all three layers.

    contract_type gates commercial scanning (NDA skips Layer 3 entirely).

    For each layer the original_text is sliced from the text the layer ran
    against (raw text for counterparty; counterparty-redacted text for
    presidio; presidio-redacted text for commercial). Counterparty matches
    are collected BEFORE substitution and applied longest-first; substitution
    is then performed in reverse char-order so recorded offsets stay valid
    against the text that the NEXT layer sees.

    Writes <stem>.review.json with confirmed=null on every span. Does NOT
    write .anon.txt. Returns the spans list.
    """
    source = pathlib.Path(source)
    output_dir = pathlib.Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    raw_text = extract_text(source)

    # B0.1: extract document-local defined-term bindings
    local_aliases = extract_defined_term_bindings(raw_text, mapping)

    # Merge local_aliases into a per-call alias map (local overrides global)
    effective_alias_map: dict[str, str] = {}
    if alias_map:
        effective_alias_map.update(alias_map)
    effective_alias_map.update(local_aliases)  # document-local wins

    # B0.7: sanitize output filename stem
    sanitized_stem = _sanitize_stem(source.stem, mapping, effective_alias_map)
    if sanitized_stem != source.stem:
        log.info('B0.7 stem sanitized: %r → %r', source.stem, sanitized_stem)

    # B0.3: compute structural ranges for exclusion
    # Ordering constraint: counterparty detection runs FIRST (below), then
    # structural ranges that contain NO counterparty span are excluded.
    structural_ranges = _compute_structural_ranges(raw_text)

    # B0.4: signature-block zones — suppress LOCATION/NRP/DATE_TIME inside them
    signature_zones = _compute_signature_zones(raw_text)
    if signature_zones:
        log.info('B0.4 signature zones detected: %d', len(signature_zones))

    spans: list[dict] = []

    # ── Layer 1a: full counterparty names (raw_text coordinates) ───────────
    # Full names run against raw_text — they have nothing to hide from yet.
    cp_full_matches: list[dict] = []
    ordered = sorted(mapping.items(), key=lambda kv: -len(kv[0]))
    for name, token in ordered:
        # Defect 1: word-boundary guard so short map names (e.g. ASE/ACE/EDE)
        # cannot match inside common words (purch[ase], pl[ace]). The
        # lookbehind/lookahead form is \b-equivalent for alphanumeric edges and
        # safe for names with leading/trailing punctuation.
        pattern = re.compile(r'(?<!\w)' + re.escape(name) + r'(?!\w)', re.IGNORECASE)
        for match in pattern.finditer(raw_text):
            cp_full_matches.append({
                'layer': 'counterparty',
                'entity_type': token,
                'token': token,
                'char_start': match.start(),
                'char_end': match.end(),
                'original_text': raw_text[match.start():match.end()],
                'score': 1.0,
            })

    # B0.3 ordering constraint (raw_text coords): exclude structural ranges
    # that do NOT contain a counterparty full-name span. A structural line with
    # a counterparty match stays.
    cp_match_positions: set[int] = {
        i for m in cp_full_matches
        for i in range(m['char_start'], m['char_end'])
    }
    cp_full_matches = [
        m for m in cp_full_matches
        if not any(
            s <= m['char_start'] < e and not cp_match_positions.intersection(range(s, e))
            for s, e in structural_ranges
        )
    ]

    # Drop spans fully contained within a longer span (e.g. 'CEO' inside
    # 'Co-Founder and CEO'). Overlapping/contained matches corrupt the offset
    # application below, producing glued sub-word tokens. (tuning 2026-06-20)
    _kept: list[dict] = []
    for m in sorted(cp_full_matches, key=lambda d: -(d['char_end'] - d['char_start'])):
        if any(o['char_start'] <= m['char_start'] and m['char_end'] <= o['char_end']
               for o in _kept):
            continue
        _kept.append(m)
    cp_full_matches = sorted(_kept, key=lambda d: d['char_start'])

    # Build cp_redacted from the full-name pass (reverse char-order, raw_text)
    # so its offsets remain valid for the alias pass that runs against it.
    cp_redacted = raw_text
    for m in sorted(cp_full_matches, key=lambda d: d['char_start'], reverse=True):
        cp_redacted = (
            cp_redacted[:m['char_start']] + m['token'] + cp_redacted[m['char_end']:]
        )

    # ── Layer 1b: alias variants (FAIL-1 fix — run on cp_redacted) ─────────
    # Running the alias loop on cp_redacted (not raw_text) prevents alias
    # fragments (e.g. "Williams") from matching inside an already-redacted
    # PARTY-XXXX token or inside the full company name before it was replaced.
    # Offsets are cp_redacted coordinates; apply_file applies full names first,
    # then aliases, to keep these offsets valid.
    alias_matches: list[dict] = []
    if effective_alias_map:
        alias_ordered = sorted(effective_alias_map.items(), key=lambda kv: -len(kv[0]))
        for alias, token in alias_ordered:
            if len(alias) < 3:
                continue
            if alias.lower() in ALIAS_STOPWORDS and alias not in local_aliases:
                continue
            pattern = re.compile(r'\b' + re.escape(alias) + r'\b', re.IGNORECASE)
            for match in pattern.finditer(cp_redacted):
                alias_matches.append({
                    'layer': 'counterparty_alias',
                    'entity_type': token,
                    'token': token,
                    'char_start': match.start(),
                    'char_end': match.end(),
                    'original_text': cp_redacted[match.start():match.end()],
                    'score': 1.0,
                })

    # Build alias_redacted from cp_redacted (reverse char-order, cp_redacted
    # coords) — this is the text the presidio layer runs against.
    alias_redacted = cp_redacted
    for m in sorted(alias_matches, key=lambda d: d['char_start'], reverse=True):
        alias_redacted = (
            alias_redacted[:m['char_start']] + m['token'] + alias_redacted[m['char_end']:]
        )

    # Order recorded spans by char position for stable, human-readable review.
    cp_full_matches.sort(key=lambda d: d['char_start'])
    alias_matches.sort(key=lambda d: d['char_start'])
    spans.extend(cp_full_matches)
    spans.extend(alias_matches)
    # Combined list retained only for the summary-log count below.
    cp_matches = cp_full_matches + alias_matches

    # ── Layer 2: presidio (runs on the alias-redacted text) ────────────────
    analyzer, _ = _presidio_engines()
    presidio_results = analyzer.analyze(
        text=alias_redacted, entities=PRESIDIO_ENTITIES, language='en',
    )
    # B0.3: drop presidio spans lying entirely within a structural range.
    # B0.4: also suppress LOCATION/NRP/DATE_TIME spans inside signature zones.
    # tuning 2026-06-20 (offset-coordinate fix): presidio_results offsets index
    # alias_redacted; structural_ranges/signature_zones were computed on raw_text.
    # Body masking diverges the two, wrongly dropping real spans (e.g. the
    # 'four (4) weeks' timeline DATE_TIME). Recompute on alias_redacted.
    _pii_structural = _compute_structural_ranges(alias_redacted)
    _pii_sigzones = _compute_signature_zones(alias_redacted)
    presidio_results = [
        r for r in presidio_results
        if not any(s <= r.start and r.end <= e for s, e in _pii_structural)
        and not (
            r.entity_type in SIGNATURE_ZONE_SUPPRESS
            and any(s <= r.start and r.end <= e for s, e in _pii_sigzones)
        )
        # tuning 2026-06-20: drop generic-word / role-label / heading FPs
        and alias_redacted[r.start:r.end].strip().lower() not in PRESIDIO_STOPWORDS
    ]
    pii_matches: list[dict] = []
    for r in presidio_results:
        pii_matches.append({
            'layer': 'presidio',
            'entity_type': r.entity_type,
            'token': f'[{r.entity_type}]',
            'char_start': r.start,
            'char_end': r.end,
            'original_text': alias_redacted[r.start:r.end],
            'score': round(r.score, 3),
        })
    pii_matches.sort(key=lambda d: d['char_start'])
    spans.extend(pii_matches)

    # Apply presidio on alias_redacted to produce the input for the commercial layer.
    _, anonymizer = _presidio_engines()
    operators = {
        entity: OperatorConfig('replace', {'new_value': f'[{entity}]'})
        for entity in PRESIDIO_ENTITIES
    }
    presidio_redacted = anonymizer.anonymize(
        text=alias_redacted,
        analyzer_results=presidio_results,
        operators=operators,
    ).text

    # B0.6: skip commercial layer for NDA contract types
    skip_commercial = contract_type.strip().lower() in CONTRACT_TYPE_SKIP_COMMERCIAL
    if skip_commercial:
        log.info('B0.6 contract_type=%r → commercial layer skipped', contract_type)

    # ── Layer 3: commercial regex (runs on presidio-redacted text) ─────────
    # com_matches initialised unconditionally so the summary log (and any later
    # reference) stays bound even when the commercial layer is skipped.
    com_matches: list[dict] = []
    if not skip_commercial:
        for pattern, token in COMMERCIAL_PATTERNS:
            for match in pattern.finditer(presidio_redacted):
                com_matches.append({
                    'layer': 'commercial_regex',
                    'entity_type': token,
                    'token': token,
                    'char_start': match.start(),
                    'char_end': match.end(),
                    'original_text': presidio_redacted[match.start():match.end()],
                    'score': 1.0,
                    'context_type': _commercial_context(
                        presidio_redacted, match.start(), match.end()
                    ),
                })
        # B2.1: TERM_DURATION context-gated spans
        com_matches.extend(_term_duration_spans(presidio_redacted))
        # B0.3: drop commercial spans lying entirely within a structural range.
        # Placed before spans.extend so the exclusion actually reaches the output.
        # tuning 2026-06-20 (offset-coordinate fix): com_matches offsets are in
        # presidio_redacted coords; structural_ranges were computed on raw_text.
        # As body masking grows the two diverge, wrongly excluding real commercial
        # terms (e.g. 'four (4) weeks', table '$14,000.00'). Recompute structural
        # ranges on presidio_redacted so the exclusion aligns to the same text.
        _com_structural = _compute_structural_ranges(presidio_redacted)
        com_matches = [
            m for m in com_matches
            if not any(s <= m['char_start'] and m['char_end'] <= e for s, e in _com_structural)
        ]
        com_matches.sort(key=lambda d: d['char_start'])
        spans.extend(com_matches)

    # ── Assign sequential ids and finalise span schema ─────────────────────
    review_spans: list[dict] = []
    for idx, m in enumerate(spans):
        review_spans.append({
            'id': idx,
            'layer': m['layer'],
            'entity_type': m['entity_type'],
            'original_text': m['original_text'],
            'proposed_token': m['token'],
            'char_start': m['char_start'],
            'char_end': m['char_end'],
            'score': m['score'],
            'score_tier': _score_tier(m['layer'], m['entity_type'], m['score']),
            'context_type': m.get('context_type'),
            'confirmed': None,
        })

    # ── Phase E: mandatory human-review triggers ───────────────────────────
    # Evaluated after spans are fully assembled, before return. Each trigger is
    # {code, severity (HARD_STOP|FLAG|HIGHLIGHT), message, span_indices}.
    # span_indices are positions in the review_spans array (== span 'id').
    ctx = context or {}
    review_triggers: list[dict] = []

    # E1 — zero counterparty redactions with a known counterparty (possible miss)
    cp_span_count = len([s for s in review_spans if s['layer'] == 'counterparty'])
    if cp_span_count == 0 and (ctx.get('party_2', '') or '').strip() != '':
        review_triggers.append({
            'code': 'E1_NO_COUNTERPARTY_SPANS',
            'severity': 'FLAG',
            'message': '0 counterparty spans detected but party_2 is set — possible detection miss',
            'span_indices': [],
        })

    # E2 — Presidio PERSON entities in the 0.35–0.60 confidence band
    e2_indices = [
        i for i, s in enumerate(review_spans)
        if s['entity_type'] == 'PERSON' and 0.35 <= s['score'] < 0.60
    ]
    if e2_indices:
        review_triggers.append({
            'code': 'E2_PERSON_LOW_CONFIDENCE',
            'severity': 'FLAG',
            'message': f'{len(e2_indices)} PERSON entity/entities in 0.35–0.60 confidence band — require explicit decision',
            'span_indices': e2_indices,
        })

    # E3 — liability_cap commercial spans (default confirmed but flagged for verify)
    e3_indices = [
        i for i, s in enumerate(review_spans)
        if s.get('context_type') == 'liability_cap'
    ]
    if e3_indices:
        review_triggers.append({
            'code': 'E3_LIABILITY_CAP',
            'severity': 'HIGHLIGHT',
            'message': f'{len(e3_indices)} liability cap amount(s) detected — default confirmed but verify',
            'span_indices': e3_indices,
        })

    # E4 — short extraction (probable scanned/image PDF)
    if len(raw_text) < 500:
        review_triggers.append({
            'code': 'E4_SHORT_EXTRACTION',
            'severity': 'FLAG',
            'message': f'Extracted text is only {len(raw_text)} chars — possible scanned/image PDF; anonymization may be incomplete',
            'span_indices': [],
        })

    # E5 enforced by D1 HARD STOP in _run_verification()
    # E6 enforced by D3 HARD STOP in _run_verification()

    private_stem = _private_stem(source)
    payload = {
        'source_file': sanitized_stem + source.suffix,
        'source_path': str(source.resolve()),
        'generated': time.strftime('%Y-%m-%dT%H:%M:%S'),
        'contract_type': contract_type,
        'sanitized_stem': sanitized_stem,
        'private_stem': private_stem,
        'doc_output_dir': str(output_dir),
        'review_triggers': review_triggers,
        'spans': review_spans,
    }
    # W2: review.json lives in the private zone (originals + reverse-map data),
    # never in the Copilot-indexed document tree.
    review_path = review_path_for(source)
    review_path.parent.mkdir(parents=True, exist_ok=True)
    with open(review_path, 'w', encoding='utf-8') as fh:
        json.dump(payload, fh, indent=2, ensure_ascii=False)

    log.info(
        'Detected %d spans (cp=%d pii=%d com=%d) → %s',
        len(review_spans), len(cp_matches), len(pii_matches), len(com_matches),
        review_path,
    )
    return review_spans

# ---------------------------------------------------------------------------
# Apply phase
# ---------------------------------------------------------------------------

def _backup_source(source: pathlib.Path) -> pathlib.Path:
    """Write a timestamped backup of source into _anon-private\\backups\\.

    Backups live OUTSIDE the document tree (in the private zone, next to
    mapping.json / decisions_library.json) so they neither clutter the
    contract folders nor leak the original filename into shippable space.

    Backup path: <_anon-private>/backups/<source_stem>_<timestamp><ext>.bak
    Never raises — logs on failure and returns None.
    """
    try:
        ts  = time.strftime("%Y%m%d_%H%M%S")
        backup_dir = DECISIONS_LIBRARY_PATH.parent / "backups"
        backup_dir.mkdir(parents=True, exist_ok=True)
        bak = backup_dir / f"{source.stem}_{ts}{source.suffix}.bak"
        import shutil as _shutil
        _shutil.copy2(str(source), str(bak))
        log.info("D0 backup: %s → %s", source.name, bak)
        return bak
    except Exception as exc:
        log.error("D0 backup failed: %s", exc)
        return None


import re as _re_d1

def _normalize_for_scan(s: str) -> str:
    """Lowercase + collapse non-alphanumeric for fuzzy comparison."""
    return _re_d1.sub(r'[^a-z0-9]+', ' ', s.lower()).strip()


PARTY_TOKEN_PATTERN = re.compile(r'PARTY-\d{4}')


def _hit_inside_party_token(text: str, match_start: int, match_end: int) -> bool:
    """Return True if the match span falls within a PARTY-XXXX token.

    Defect 2b: a residual-name hit that lands inside an already-substituted
    PARTY token (e.g. the digits of 'PARTY-0046') is not a real leak.
    """
    for m in PARTY_TOKEN_PATTERN.finditer(text):
        if m.start() <= match_start and match_end <= m.end():
            return True
    return False


def _scan_leakage(
    anon_text: str,
    name_map:  dict[str, str],
    alias_map: dict[str, str],
    threshold: int = 90,
) -> list[dict]:
    """Scan anonymized text for residual counterparty identity leakage.

    Uses rapidfuzz token_set_ratio on normalized 6-word windows.
    Returns list of findings: {name, window, score, char_start, match_type}.
    """
    findings: list[dict] = []
    _combined_stop = {s.lower() for s in ALIAS_STOPWORDS} | {w.lower() for w in COMMON_WORDS}
    words = anon_text.split()
    candidates = {
        _normalize_for_scan(k): k
        for k in list(name_map.keys()) + list(alias_map.keys())
        if len(k) >= 4
    }

    # Slide a 6-word window across the anonymized text
    window_size = 6
    pos = 0
    for i in range(len(words)):
        window_words = words[i : i + window_size]
        window       = " ".join(window_words)
        norm_window  = _normalize_for_scan(window)
        for norm_name, orig_name in candidates.items():
            score = fuzz.token_set_ratio(norm_name, norm_window)
            if score >= threshold:
                if orig_name.lower() in _combined_stop:
                    continue
                # Drop short ALL-CAPS abbreviation aliases (e.g. CCOM, Merc, TITL)
                if len(orig_name) <= 4 and orig_name.isupper():
                    continue
                # Defect 2b: rapidfuzz gives no match offset — locate an exact
                # occurrence and drop it if it lands inside a PARTY-XXXX token.
                loc = re.search(re.escape(orig_name), anon_text, re.IGNORECASE)
                if loc and _hit_inside_party_token(anon_text, loc.start(), loc.end()):
                    continue
                findings.append({
                    "name":       orig_name,
                    "window":     window,
                    "score":      round(score, 1),
                    "char_start": pos,
                    "match_type": "fuzzy",
                })
        pos += len(window_words[0]) + 1 if window_words else 0

    # Deduplicate by name+window
    seen: set[tuple] = set()
    deduped: list[dict] = []
    for f in findings:
        key = (f["name"], f["window"])
        if key not in seen:
            seen.add(key)
            deduped.append(f)

    return deduped


def _sweep_counterparty_dict(
    anon_text: str,
    name_map:  dict[str, str],
    alias_map: dict[str, str],
    fuzzy_threshold: int = 92,
) -> list[dict]:
    """Sweep for any name_map or alias_map key that survives anonymization.

    Two passes per name:
      (a) exact — word-boundary, case-insensitive regex (match_type "exact").
      (b) fuzzy (F) — for names >= 6 chars with no exact hit, flag a survivor
          when rapidfuzz partial_ratio(name, anon_text) >= fuzzy_threshold
          (default 92, match_type "fuzzy"). partial_ratio finds the optimal
          substring alignment natively, so no manual windowing is needed.
          Tighter than D1's 90 because D2 scans the full name list, so the
          false-positive cost of a HARD STOP is higher.

    Returns list of {name, count, match_type[, score]} for each hit.
    """
    findings: list[dict] = []
    _combined_stop = {s.lower() for s in ALIAS_STOPWORDS} | {w.lower() for w in COMMON_WORDS}
    anon_lower = anon_text.lower()
    for name in sorted(
        list(name_map.keys()) + list(alias_map.keys()),
        key=len, reverse=True,
    ):
        if len(name) < 4:
            continue
        if name.lower() in _combined_stop:
            continue
        # Drop short ALL-CAPS abbreviation aliases (e.g. CCOM, Merc, TITL)
        if len(name) <= 4 and name.isupper():
            continue
        # ── (a) Exact pass (unchanged behaviour) ───────────────────────────
        # Defect 2a: word-boundary match so a dictionary name cannot flag inside
        # a longer word (e.g. "Convey" inside the un-redactable "Conveyor").
        pattern = re.compile(r'(?<!\w)' + re.escape(name) + r'(?!\w)', re.IGNORECASE)
        count = 0
        for hit in pattern.finditer(anon_text):
            # Defect 2b: ignore hits that fall inside a PARTY-XXXX token.
            if _hit_inside_party_token(anon_text, hit.start(), hit.end()):
                continue
            count += 1
        if count:
            findings.append({"name": name, "count": count, "match_type": "exact"})
            continue  # exact already caught it — fuzzy pass would be redundant

        # ── (b) Fuzzy pass (F) — names >= 6 chars, no exact hit ─────────────
        # Direct partial_ratio finds the optimal substring alignment natively,
        # so no manual window iteration is needed.
        if len(name) < 6:
            continue
        score = fuzz.partial_ratio(name.lower(), anon_lower)
        if score >= fuzzy_threshold:
            findings.append({
                "name": name, "count": 1,
                "match_type": "fuzzy", "score": round(score, 1),
            })
    return findings


def _check_roundtrip(
    review:   dict,
    name_map: dict[str, str],
) -> list[dict]:
    """Verify every confirmed counterparty span has a reverse-map entry.

    For each confirmed span whose layer is counterparty or counterparty_alias:
      - token = proposed_token (e.g. PARTY-0001)
      - reverse_map[token] → original_text must exist and match span original_text

    Returns list of {token, original_text, issue} for any mismatch.
    """
    # Build reverse map: token → original_text from confirmed spans
    reverse: dict[str, str] = {}
    for span in review.get("spans", []):
        if span.get("confirmed") and span["layer"] in (
            "counterparty", "counterparty_alias"
        ):
            tok = span["proposed_token"]
            orig = span["original_text"]
            if tok not in reverse:
                reverse[tok] = orig

    issues: list[dict] = []
    # Cross-check: each token must appear in name_map values
    name_map_values = set(name_map.values())
    for tok, orig in reverse.items():
        if tok not in name_map_values:
            issues.append({
                "token":         tok,
                "original_text": orig,
                "issue":         "token not in name_map — mapping may be stale",
            })
    return issues


def _scan_orphan_tokens(
    anon_text: str,
    name_map:  dict[str, str],
) -> list[str]:
    """Find PARTY-XXXX tokens in the anonymized text that have no reverse
    mapping in name_map.

    An orphan token means apply() substituted a token that cannot be
    de-anonymized — the reverse map is incomplete.

    Returns list of orphan token strings.
    """
    import re as _re_d4
    token_pattern = _re_d4.compile(r'\bPARTY-\d{4}\b')
    found_tokens  = set(token_pattern.findall(anon_text))
    mapped_tokens = set(name_map.values())
    return sorted(found_tokens - mapped_tokens)


def _excerpt(text: str, matched: str, before: int = 3, after: int = 3) -> str:
    """Return the line containing `matched` plus `before`/`after` lines of context."""
    if not matched:
        return ""
    lines = text.splitlines()
    low = matched.lower()
    for i, ln in enumerate(lines):
        if low in ln.lower():
            start = max(0, i - before)
            end   = min(len(lines), i + after + 1)
            return "\n".join(lines[start:end])
    return ""


def _run_verification(
    anon_text:  str,
    review:     dict,
    name_map:   dict[str, str],
    alias_map:  dict[str, str],
    stem:       str,
) -> dict:
    """Run all D-phase checks on the in-memory redacted text and write .verify.json.

    Returns the verification payload, including a unified `hits` list (each with
    check / matched_text / ±3-line excerpt) and a `shippable` flag. Output is
    shippable only when every check is clean.
    """
    leakage   = _scan_leakage(anon_text, name_map, alias_map)
    sweep     = _sweep_counterparty_dict(anon_text, name_map, alias_map)
    roundtrip = _check_roundtrip(review, name_map)
    orphans   = _scan_orphan_tokens(anon_text, name_map)

    passed = not (leakage or sweep or roundtrip or orphans)

    # Unified hit list with ±3-line excerpts for the warning response (6c).
    hits: list[dict] = []
    for f in leakage:
        hits.append({"check": "leakage", "matched_text": f["name"],
                     "excerpt": _excerpt(anon_text, f.get("window") or f["name"])})
    for f in sweep:
        hits.append({"check": "sweep", "matched_text": f["name"],
                     "excerpt": _excerpt(anon_text, f["name"])})
    for f in roundtrip:
        hits.append({"check": "roundtrip", "matched_text": f.get("token", ""),
                     "excerpt": _excerpt(anon_text, f.get("token", ""))})
    for tok in orphans:
        hits.append({"check": "orphan", "matched_text": tok,
                     "excerpt": _excerpt(anon_text, tok)})

    payload = {
        "generated":  time.strftime("%Y-%m-%dT%H:%M:%S"),
        "passed":     passed,
        "shippable":  passed,
        "leakage":    leakage,
        "sweep":      sweep,
        "roundtrip":  roundtrip,
        "orphans":    orphans,
        "hits":       hits,
        "summary": {
            "leakage_count":   len(leakage),
            "sweep_count":     len(sweep),
            "roundtrip_issues": len(roundtrip),
            "orphan_count":    len(orphans),
        },
    }

    # W2: verify.json lives in the private zone (leakage windows expose names).
    VERIFY_DIR.mkdir(parents=True, exist_ok=True)
    verify_path = VERIFY_DIR / f"{stem}.verify.json"
    try:
        verify_path.write_text(
            json.dumps(payload, indent=2, ensure_ascii=False),
            encoding="utf-8",
        )
        log.info(
            "D5 verify: passed=%s leakage=%d sweep=%d roundtrip=%d orphans=%d",
            passed, len(leakage), len(sweep), len(roundtrip), len(orphans),
        )
    except Exception as exc:
        log.error("D5 verify write failed: %s", exc)

    return payload


def apply_file(
    review_json_path: pathlib.Path,
    output_dir: pathlib.Path,
) -> dict:
    """Replay the three layers from a reviewed review.json, confirmed spans only.

    Reads review.json, reloads raw text from source_path, then for each layer
    applies ONLY spans where confirmed=true:
      - counterparty / counterparty_alias: replace in reverse char_start order
      - presidio: reconstruct RecognizerResult objects, run AnonymizerEngine
      - commercial: replace in reverse char_start order

    On verification pass, writes <stem>.anon.txt to output_dir (the document
    tree) and <stem>.audit.json to the private audits zone (confirmed spans
    only). Returns a summary dict.
    """
    review_json_path = pathlib.Path(review_json_path)
    output_dir = pathlib.Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    with open(review_json_path, encoding='utf-8') as fh:
        review = json.load(fh)

    source = pathlib.Path(review['source_path'])
    spans = review['spans']

    # D0: backup source before apply
    _backup_source(source)

    # B0.7: use sanitized stem for output filenames if available
    sanitized_stem = review.get('sanitized_stem') or source.stem
    if sanitized_stem != source.stem:
        log.info('B0.7 applying sanitized stem: %r', sanitized_stem)

    confirmed_spans = [s for s in spans if s.get('confirmed') is True]
    rejected_count = sum(1 for s in spans if s.get('confirmed') is not True)

    raw_text = extract_text(source)
    text = raw_text

    # ── Layer 1a: full counterparty names, reverse char_start order (raw) ──
    cp_full_confirmed = [s for s in confirmed_spans if s['layer'] == 'counterparty']
    for s in sorted(cp_full_confirmed, key=lambda d: d['char_start'], reverse=True):
        text = text[:s['char_start']] + s['proposed_token'] + text[s['char_end']:]

    # ── Layer 1b: aliases, reverse char_start order (cp_redacted coords) ───
    # Detection ran aliases against the full-name-redacted text, so their
    # offsets are valid only after Layer 1a has been applied.
    cp_alias_confirmed = [s for s in confirmed_spans if s['layer'] == 'counterparty_alias']
    for s in sorted(cp_alias_confirmed, key=lambda d: d['char_start'], reverse=True):
        text = text[:s['char_start']] + s['proposed_token'] + text[s['char_end']:]

    # ── Layer 2: presidio via AnonymizerEngine on the cp-redacted text ─────
    pii_confirmed = [s for s in confirmed_spans if s['layer'] == 'presidio']
    if pii_confirmed:
        _, anonymizer = _presidio_engines()
        analyzer_results = [
            RecognizerResult(
                entity_type=s['entity_type'],
                start=s['char_start'],
                end=s['char_end'],
                score=float(s.get('score') or 1.0),
            )
            for s in pii_confirmed
        ]
        operators = {
            s['entity_type']: OperatorConfig(
                'replace', {'new_value': s['proposed_token']}
            )
            for s in pii_confirmed
        }
        text = anonymizer.anonymize(
            text=text,
            analyzer_results=analyzer_results,
            operators=operators,
        ).text

    # ── Layer 3: commercial regex, reverse char_start order ────────────────
    com_confirmed = [s for s in confirmed_spans if s['layer'] == 'commercial_regex']
    for s in sorted(com_confirmed, key=lambda d: d['char_start'], reverse=True):
        text = text[:s['char_start']] + s['proposed_token'] + text[s['char_end']:]

    # ── FAIL-2: verify the in-memory redacted text BEFORE writing anything ──
    # Nothing reaches the document tree until verification passes. A failing
    # run is quarantined in the private zone, never shipped.
    private_stem = review.get('private_stem') or _private_stem(source)
    name_map, alias_map = get_mapping(
        PRIVATE_DIR / "mapping.json",
        pathlib.Path(__file__).resolve().parent.parent / "kb" / "COUNTERPARTIES.md",
        rebuild=False,
    )
    verification = _run_verification(
        anon_text = text,          # in-memory, not a written file
        review    = review,
        name_map  = name_map,
        alias_map = alias_map,
        stem      = private_stem,
    )

    cp_reds  = [s for s in confirmed_spans if s['layer'] in ('counterparty', 'counterparty_alias')]
    pii_reds = [s for s in confirmed_spans if s['layer'] == 'presidio']
    com_reds = [s for s in confirmed_spans if s['layer'] == 'commercial_regex']

    if not verification["passed"]:
        # Quarantine OUTSIDE the document tree; do NOT write .anon.txt/.audit.json there.
        QUARANTINE_DIR.mkdir(parents=True, exist_ok=True)
        q_path = quarantine_path_for(source)
        q_path.write_text(text, encoding="utf-8")
        log.warning(
            "VERIFY FAILED — quarantined %s (leakage=%d sweep=%d roundtrip=%d orphans=%d)",
            q_path.name,
            verification["summary"]["leakage_count"],
            verification["summary"]["sweep_count"],
            verification["summary"]["roundtrip_issues"],
            verification["summary"]["orphan_count"],
        )
        return {
            'passed':          False,
            'quarantine_path': str(q_path),
            'hits':            verification.get("hits", []),
            'confirmed':       len(confirmed_spans),
            'rejected':        rejected_count,
            'verify':          verification,
        }

    # ── Passed: write shippable output to the document tree ─────────────────
    anon_path = output_dir / f'{sanitized_stem}.anon.txt'
    anon_path.write_text(text, encoding='utf-8')

    audit_payload = {
        'source_file': sanitized_stem + source.suffix,
        'private_stem': private_stem,
        'layers_applied': ['counterparty', 'presidio', 'commercial_regex'],
        'redaction_count': {
            'counterparty': len(cp_reds),
            'presidio':     len(pii_reds),
            'commercial_regex': len(com_reds),
        },
        'redactions':       cp_reds,
        'presidio_entities': pii_reds,
        'commercial_terms': com_reds,
    }
    # audit.json holds original_text↔token reverse-map data → private zone,
    # NOT the Copilot-indexed document tree.
    audit_path = audit_path_for(output_dir, sanitized_stem)
    audit_path.parent.mkdir(parents=True, exist_ok=True)
    with open(audit_path, 'w', encoding='utf-8') as fh:
        json.dump(audit_payload, fh, indent=2, ensure_ascii=False)

    log.info(
        'Applied %d confirmed spans (%d rejected) → %s',
        len(confirmed_spans), rejected_count, anon_path,
    )

    return {
        'passed':     True,
        'anon_path':  str(anon_path),
        'audit_path': str(audit_path),
        'confirmed':  len(confirmed_spans),
        'rejected':   rejected_count,
        'verify':     verification,
    }

# ---------------------------------------------------------------------------
# Pipeline orchestration (CLI --full-run)
# ---------------------------------------------------------------------------

def process_file(
    source: pathlib.Path,
    mapping: dict[str, str],
    output_dir: pathlib.Path,
    write_audit: bool,
    alias_map: Optional[dict[str, str]] = None,
) -> dict:
    """Full pipeline: detect, auto-confirm every span, then apply.

    Preserves the original CLI behaviour (all detected spans are redacted)
    while routing through the new detect/apply machinery. Never raises —
    errors are captured in the summary.
    """
    t0 = time.monotonic()
    log.info('Processing: %s', source.name)

    try:
        spans = detect_file(source, mapping, alias_map, output_dir)
    except Exception as exc:
        log.error('Detection failed — %s: %s', source.name, exc)
        return {
            'file': str(source),
            'status': 'extraction_error',
            'error': str(exc),
        }

    # Auto-confirm every detected span in the review.json, then apply.
    review_path = review_path_for(source)
    try:
        with open(review_path, encoding='utf-8') as fh:
            review = json.load(fh)
        for s in review['spans']:
            s['confirmed'] = True
        with open(review_path, 'w', encoding='utf-8') as fh:
            json.dump(review, fh, indent=2, ensure_ascii=False)

        result = apply_file(review_path, output_dir)
    except Exception as exc:
        log.error('Redaction failed — %s: %s', source.name, exc)
        return {
            'file': str(source),
            'status': 'redaction_error',
            'error': str(exc),
        }

    # FAIL-2: a failed-verification apply returns no doc-tree output — it was
    # quarantined. Report it as such instead of assuming an audit_path exists.
    if not result.get('passed', False):
        elapsed = time.monotonic() - t0
        log.warning('  %-50s  QUARANTINED (verification failed) | %.1fs',
                    source.name, elapsed)
        return {
            'file': source.name,
            'status': 'quarantined',
            'quarantine_path': result.get('quarantine_path'),
            'verify': result.get('verify'),
            'runtime_s': round(elapsed, 2),
        }

    if not write_audit:
        # process_file historically wrote audit only on --audit; honour that.
        audit_path = pathlib.Path(result['audit_path'])
        if audit_path.exists():
            audit_path.unlink()

    cp_n = sum(1 for s in spans if s['layer'] in ('counterparty', 'counterparty_alias'))
    pii_n = sum(1 for s in spans if s['layer'] == 'presidio')
    com_n = sum(1 for s in spans if s['layer'] == 'commercial_regex')

    elapsed = time.monotonic() - t0
    summary = {
        'file': source.name,
        'status': 'ok',
        'counterparty_redactions': cp_n,
        'presidio_redactions':     pii_n,
        'commercial_redactions':   com_n,
        'runtime_s': round(elapsed, 2),
    }
    log.info(
        '  %-50s  cp=%-3d pii=%-3d com=%-3d | %.1fs',
        source.name, cp_n, pii_n, com_n, elapsed,
    )
    return summary

# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        prog='anonymize.py',
        description=(
            'P3 contract anonymization pipeline. '
            'Applies counterparty, PII, and commercial-term redaction in three ordered layers.'
        ),
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            'Examples:\n'
            '  python anonymize.py --input contract.docx --detect-only\n'
            '  python anonymize.py --input contract.pdf --full-run --audit\n'
            '  python anonymize.py --input contracts/ --output-dir anonymization/output --audit\n'
            '  python anonymize.py --input contract.pdf --rebuild-map\n'
        ),
    )
    parser.add_argument(
        '--input', required=True, metavar='PATH',
        help='Source file (.pdf/.docx/.doc/.txt) or directory of source files.',
    )
    parser.add_argument(
        '--output-dir', default='anonymization/output', metavar='DIR',
        help='Directory for output. (default: anonymization/output)',
    )
    parser.add_argument(
        '--mapping',
        default=r'C:\Users\jrudy\OneDrive - Diakonia Group, LLC\Contract Management - SharePoint\_anon-private\mapping.json',
        metavar='FILE',
        help='Counterparty mapping JSON (sensitive — stored outside the git tree in _anon-private\\).',
    )
    parser.add_argument(
        '--counterparties', default='kb/COUNTERPARTIES.md', metavar='FILE',
        help='Source for mapping construction. (default: kb/COUNTERPARTIES.md)',
    )
    parser.add_argument(
        '--audit', action='store_true',
        help='Write a .audit.json redaction log alongside each .anon.txt output (full-run only).',
    )
    parser.add_argument(
        '--detect-only', action='store_true',
        help='Run detection only: write <stem>.review.json and print its path. No .anon.txt.',
    )
    parser.add_argument(
        '--full-run', action='store_true',
        help='Detect, auto-confirm all spans, and apply (legacy one-shot behaviour).',
    )
    parser.add_argument(
        '--rebuild-map', action='store_true',
        help='Force rebuild of mapping.json from --counterparties source.',
    )
    return parser


def main() -> None:
    parser = build_arg_parser()
    args = parser.parse_args()

    output_dir       = pathlib.Path(args.output_dir)
    mapping_path     = pathlib.Path(args.mapping)
    counterparties_p = pathlib.Path(args.counterparties)

    output_dir.mkdir(parents=True, exist_ok=True)

    mapping, alias_map = get_mapping(mapping_path, counterparties_p, args.rebuild_map)

    input_path = pathlib.Path(args.input)
    if input_path.is_file():
        files = [input_path]
    elif input_path.is_dir():
        files = sorted(
            p for p in input_path.rglob('*')
            if p.suffix.lower() in SUPPORTED_EXTENSIONS
        )
    else:
        parser.error(f'--input path does not exist: {input_path}')

    log.info('Files to process: %d', len(files))

    if args.detect_only:
        for file_path in files:
            try:
                detect_file(file_path, mapping, alias_map, output_dir)
                review_path = review_path_for(file_path)
                print(review_path)
            except Exception as exc:
                log.error('Detection failed — %s: %s', file_path.name, exc)
        return

    summaries: list[dict] = []
    for file_path in files:
        summary = process_file(file_path, mapping, output_dir, args.audit, alias_map)
        summaries.append(summary)

    ok_count    = sum(1 for s in summaries if s.get('status') == 'ok')
    error_count = len(summaries) - ok_count
    log.info('Finished: %d succeeded, %d failed.', ok_count, error_count)

    if error_count:
        for s in summaries:
            if s.get('status') != 'ok':
                log.error('  %s: %s — %s', s['file'], s['status'], s.get('error', ''))


if __name__ == '__main__':
    main()
