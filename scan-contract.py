"""
scan-contract.py  --  Contract Catalog Scanner
Reads contract files (.pdf, .docx), extracts metadata, and updates contract-catalog.csv.

Usage (from SharePoint root):
  python Tools/scan-contract.py "VendorFolder/filename.pdf"
  python Tools/scan-contract.py --vendor "Disney"
  python Tools/scan-contract.py --location "02 Unsigned Contracts"
  python Tools/scan-contract.py --all [--dry-run]
"""

import argparse
import csv
import re
import shutil
import sys
import time
from datetime import date
from pathlib import Path, PurePosixPath

import pandas as pd
from dateutil import parser as dparser
from dateutil.relativedelta import relativedelta

# ── Paths ─────────────────────────────────────────────────────────────────────

SCRIPT_DIR   = Path(__file__).resolve().parent          # …/Tools
SHAREPOINT   = SCRIPT_DIR.parent                        # …/Contract Management - SharePoint
ONEDRIVE     = SHAREPOINT.parent                        # …/OneDrive - Diakonia Group, LLC
DEFAULT_CSV  = SCRIPT_DIR / "contract-catalog.csv"

# ContractLocation is a LOGICAL label, not always a folder name. "01 Active
# Contracts" has no folder by that name: active files physically live in the
# separately-synced "Salesforce Integration - Active Contracts" library, a
# sibling of the SharePoint root. Always resolve through LOCATION_ROOTS —
# never join SHAREPOINT / ContractLocation directly.
LOCATION_ROOTS = {
    "01 Active Contracts":   ONEDRIVE   / "Salesforce Integration - Active Contracts",
    "02 Unsigned Contracts": SHAREPOINT / "02 Unsigned Contracts",
    "03 Archived Contracts": SHAREPOINT / "03 Archived Contracts",
    "04 Expired Contracts":  SHAREPOINT / "04 Expired Contracts",
}

CONTRACT_ROOTS = list(LOCATION_ROOTS)

# ── DCS identity ──────────────────────────────────────────────────────────────

DCS_RE = re.compile(
    r"Designed\s+Conveyor\s+Systems(?:,?\s+LLC)?"
    r"|\bDCS\b",
    re.IGNORECASE,
)

# ── Doc-type patterns (order matters; checked against first 500 chars) ────────

DOC_TYPE_PATTERNS = [
    (re.compile(r"MUTUAL\s+NON.?DISCLOSURE",                         re.I), "MNDA"),
    (re.compile(r"MUTUAL\s+CONFIDENTIALITY\s+AND\s+NON.?DISCLOSURE", re.I), "MNDA"),
    (re.compile(r"NON.?DISCLOSURE",                                   re.I), "NDA"),
    (re.compile(r"MASTER\s+SERVICES?\s+AGREEMENT",                   re.I), "MSA"),
    (re.compile(r"PROFESSIONAL\s+SERVICES?\s+AGREEMENT",             re.I), "PSA"),
    (re.compile(r"MASTER\s+PURCHASING\s+AGREEMENT",                  re.I), "MPA"),
    (re.compile(r"INTEGRATOR\s+AGREEMENT|PARTNER\s+AGREEMENT"
                r"|DISTRIBUTOR\s+AGREEMENT",                          re.I), "IPA"),
    (re.compile(r"STATEMENT\s+OF\s+WORK",                            re.I), "SOW"),
    (re.compile(r"\bPURCHASE\s+ORDER\b|\bPROCUREMENT\s+ORDER\b"
                r"|\bCONTRACT\s+AWARD\b|\bDELIVERY\s+ORDER\b",       re.I), "PO"),
    (re.compile(r"SUBCONTRACT(?!OR)",                                 re.I), "Subcontract"),
    (re.compile(r"CHANGE\s+ORDER",                                    re.I), "Change Order"),
    (re.compile(r"NON.?COMPETE|NONCOMPETE",                          re.I), "Non-Compete"),
    (re.compile(r"\bEND.USER\s+LICENSE\s+AGREEMENT\b|\bEULA\b",      re.I), "EULA"),
    (re.compile(r"\bLICENSE\s+AGREEMENT\b",                          re.I), "License"),
    (re.compile(r"\bEXHIBIT\b",                                      re.I), "Exhibit"),
    (re.compile(r"\bAMENDMENT\b",                                    re.I), "Amendment"),
    (re.compile(r"\bRETAINER\b",                                     re.I), "Retainer"),
    (re.compile(r"\bFORM\s+W-?9\b|\bW-9\b"
                r"|REQUEST\s+FOR\s+TAXPAYER\s+IDENTIFICATION",        re.I), "Tax-Form"),
]

# DocTypes that are not bilateral contracts — skip signing detection for these
_NON_CONTRACT_TYPES = {"Tax-Form", "Exhibit"}

# ── Signature patterns ────────────────────────────────────────────────────────

# DocuSign: Envelope ID in contract body OR Certificate Of Completion page
DOCUSIGN_DOC_RE = re.compile(
    r"DocuSign\s+Envelope\s+(?:ID|Id)\b"
    r"|Certificate\s+Of\s+Completion\b",
    re.IGNORECASE,
)

# DocuSign signer timestamp rendered at end of doc: "10/31/2024 | 10:52 AM PDT"
DOCUSIGN_TIMESTAMP_RE = re.compile(
    r"\d{1,2}/\d{1,2}/\d{4}\s*\|\s*\d{1,2}:\d{2}\s*(?:AM|PM)",
    re.IGNORECASE,
)

# E-signature platform audit trail indicators
ADOBE_AUDIT_RE = re.compile(
    r"Transaction\s+ID[:\s]"                   # Adobe Sign audit field
    r"|Document\s+e-signed\s+by\b"             # Adobe Sign event
    r"|E-SIGNED\s+by\b"                         # USPS / custom e-sign
    r"|Agreement\s+completed\."                 # Adobe Sign completion
    r"|Adobe\s*(?:Acrobat\s*)?Sign\b"           # Adobe Sign branding
    r"|Signed\s+with\s+PandaDoc\b"              # PandaDoc
    r"|Signer\s+ID:",                           # PandaDoc audit field
    re.IGNORECASE,
)

# E-signature completion signals
ADOBE_COMPLETED_RE = re.compile(
    r"Agreement\s+completed\."
    r"|Status[:\s]+\n?Signed\b"
    r"|Signed\s+with\s+PandaDoc\b",            # PandaDoc completion marker
    re.IGNORECASE,
)

# Any "e-signed by" variant (with or without "Document" prefix)
ESIGNED_BY_RE = re.compile(r"(?:Document\s+)?[Ee]-[Ss]igned\s+by\b", re.IGNORECASE)

# Adobe Sign / PandaDoc certificate: "Name (Month D, YYYY HH:MM TZ)"
ESIG_CERT_RE = re.compile(
    r"[A-Z][a-zA-Z\-']+(?:\s+[A-Z][a-zA-Z\-']+)+\s*"
    r"\([A-Za-z]{3}\w*\s+\d{1,2},?\s+\d{4}\s+\d{1,2}:\d{2}",
    re.IGNORECASE,
)

# Electronic signature marker: /s/ Name
SLASH_S_RE = re.compile(r"/s/\s*.{2,80}?(?:\n|$)")

PLACEHOLDER_RE = re.compile(
    r"^[\s_\-\*]*$"
    r"|^NAME$"
    r"|\(VENDOR"
    r"|_{3,}"
    r"|^\[",
    re.IGNORECASE,
)

# DocuSign template field placeholders: \s1\, \fullname1\, $docusign:SignHere::, etc.
DOCUSIGN_FIELD_RE = re.compile(r"\\[a-zA-Z]\w*\\|\$docusign:[^\s]+")

# Broad date pattern for DocuSign tail: handles DD-Mon-YYYY, M/D/YYYY, Month D YYYY, etc.
_MON3 = r"Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec"
DATE_ANY_RE = re.compile(
    r"\b\d{1,2}[-/](?:\d{1,2}|" + _MON3 + r")[-/]\d{2,4}\b"
    r"|\b(?:" + _MON3 + r")[a-z]*\s+\d{1,2},?\s+\d{4}\b"
    r"|\b\d{1,2}\s+(?:" + _MON3 + r")[a-z]*\s+\d{4}\b",
    re.IGNORECASE,
)


# ── Date patterns ─────────────────────────────────────────────────────────────

_MONTH = (
    r"(?:January|February|March|April|May|June|July|August|"
    r"September|October|November|December)"
)
DATE_PAT = (
    r"(?:"
    + _MONTH + r"\s+\d{1,2},?\s+\d{4}"          # October 1, 2024
    r"|\d{1,2}\s+" + _MONTH + r"\s+\d{4}"        # 1 October 2024
    r"|\d{4}-\d{2}-\d{2}"                         # 2024-10-01
    r"|\d{1,2}/\d{1,2}/\d{2,4}"                   # 10/1/2024
    r"|\d{1,2}\.\d{1,2}\.\d{4}"                   # 10.1.2024
    r")"
)

# Date label in a signature block — used as EffectiveDate fallback
SIG_DATE_RE = re.compile(
    r"(?:Date(?:\s+Signed)?|Dated?)\s*[:\s]\s*(" + DATE_PAT + r")",
    re.IGNORECASE,
)

EFFECTIVE_DATE_RE = re.compile(
    r"(?:effective\s+(?:as\s+of\s+)?|dated\s+(?:as\s+of\s+)?"
    r"|as\s+of\s+|entered\s+into\s+(?:as\s+of\s+)?)"
    + DATE_PAT,
    re.IGNORECASE,
)

# Handles fill-in-blank templates: "made as of the _30_ day of _April_____, 2020"
_AS_OF_THE_RE = re.compile(
    r"(?:effective|dated|made)\s+as\s+of\s+the\s+|as\s+of\s+the\s+",
    re.IGNORECASE,
)
_DAY_OF_RE = re.compile(
    r"_*(\d{1,2})_*\s+day\s+of\s+_*("
    + _MONTH[3:]          # strip leading "(?:" → becomes capturing group (Month|...)
    + r"_*[_,\s]+_*(\d{4})",
    re.IGNORECASE,
)

# At-will / written-notice termination (no fixed expiration date)
AT_WILL_RE = re.compile(
    r"either\s+party\s+may\s+terminat"
    r"|terminat\s+this\s+agreement\s+at\s+any\s+time"
    r"|upon\s+(?:thirty|sixty|ninety|\d+)\s+days?\s+(?:written\s+)?notice"
    r"|upon\s+written\s+notice\s+to\s+the\s+other",
    re.IGNORECASE,
)

EXPIRATION_TRIGGER_RE = re.compile(
    r"(?:expir(?:es?|ation)|terminat(?:es?|ion\s+date)|through|until|ending\s+on)"
    r"\s+(?:on\s+)?"
    + DATE_PAT,
    re.IGNORECASE,
)

TERM_YEARS_RE = re.compile(
    r"(?:term\s+of\s+|for\s+a\s+(?:period\s+of\s+)?)"
    r"(\w+|\d+)\s*\(?\d*\)?\s*year(?:s)?",
    re.IGNORECASE,
)

WORD_TO_NUM = {
    "one": 1, "two": 2, "three": 3, "four": 4, "five": 5,
    "six": 6, "seven": 7, "eight": 8, "nine": 9, "ten": 10,
}

# ── Entity suffix for counterparty fallback ───────────────────────────────────

ENTITY_SUFFIX_RE = re.compile(
    r"\b(?:Inc\.?|LLC\.?|Corp\.?|Ltd\.?|L\.P\.?|LLP\.?|Co\.?|plc|GmbH|B\.V\.|S\.A\.?)\b",
    re.IGNORECASE,
)

# Specific: use ", and" (comma + and) as party separator — handles names containing "and"
PARTY_COMMA_AND_RE = re.compile(
    r"(?:by\s+and\s+between|between)\s+"
    r"(.{10,350}?)"
    r",\s+and\s+"
    r"(.{5,300}?)"
    r"(?:,\s*(?:\(|located|having|a\s+\w)|[\(\.\n\"])",
    re.IGNORECASE | re.DOTALL,
)

# Generic fallback: plain "and" separator
PARTY_BLOCK_RE = re.compile(
    r"(?:by\s+and\s+between|between)"
    r"\s+(.{5,200}?)"
    r"\s+and\s+"
    r"(.{5,200}?)"
    r"(?:\s*[\(\.\"]|,\s*(?:each|a\s+.Party.|together))",
    re.IGNORECASE | re.DOTALL,
)

# DCS boilerplate: "DCS and its Affiliates ..., on the one hand, and [COUNTERPARTY]"
DCS_ONE_HAND_RE = re.compile(
    r"(?:Designed\s+Conveyor\s+Systems[^,\n]*|DCS)[^,\n]*"
    r",\s*on\s+the\s+one\s+hand,\s+and\s+"
    r"(.{5,200}?)"
    r"(?:\s*[\(\"\'\.]|\s*,\s*(?:on\s+the\s+other\s+hand|each|together|\"|\'))",
    re.IGNORECASE | re.DOTALL,
)

# "and [ENTITY], on the other hand"
OTHER_HAND_RE = re.compile(
    r"and\s+(.{5,200}?),\s*on\s+the\s+other\s+hand",
    re.IGNORECASE | re.DOTALL,
)

# ── Text extraction ───────────────────────────────────────────────────────────

def _extract_pdf(path: Path) -> tuple[str, bool]:
    import fitz  # PyMuPDF
    try:
        doc = fitz.open(str(path))
    except PermissionError:
        raise
    except Exception as exc:
        raise RuntimeError(f"fitz could not open {path.name}: {exc}") from exc
    pages = [page.get_text() for page in doc]
    full = "\n".join(pages)
    return full, len(full.strip()) < 50


def _extract_docx(path: Path) -> tuple[str, bool]:
    from docx import Document
    try:
        doc = Document(str(path))
    except PermissionError:
        raise
    except Exception as exc:
        raise RuntimeError(f"python-docx could not open {path.name}: {exc}") from exc
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts), False


def extract_text(path: Path) -> tuple[str, bool, str | None]:
    """Returns (text, is_scanned, skip_reason)."""
    if path.name.startswith("~$"):
        return ("", False, "Word temp file (~$) — excluded")
    ext = path.suffix.lower()
    if ext == ".pdf":
        try:
            return (*_extract_pdf(path), None)
        except PermissionError:
            return ("", False, "file locked (PermissionError)")
        except RuntimeError as e:
            return ("", False, str(e))
    elif ext == ".docx":
        try:
            return (*_extract_docx(path), None)
        except PermissionError:
            return ("", False, "file locked (PermissionError)")
        except RuntimeError as e:
            return ("", False, str(e))
    elif ext == ".doc":
        return ("", False, ".doc not supported (legacy binary format)")
    elif ext == ".msg":
        return ("", False, ".msg not supported")
    else:
        return ("", False, f"{ext} not supported")

# ── Extractors ────────────────────────────────────────────────────────────────

def detect_doc_type(text: str) -> str | None:
    header = text[:500]
    for pattern, doc_type in DOC_TYPE_PATTERNS:
        if pattern.search(header):
            return doc_type
    return None


def _count_filled_by_blocks(text: str) -> int:
    """
    Count By: signature blocks where a real name appears within 400 chars.
    Handles names on the same line, next line, or several blank lines below.
    Skips placeholders, DocuSign field tags, form labels, and contract body text.
    """
    count = 0
    skip_labels = re.compile(
        r"^(?:Date|Name|Title|Email|Signature|Print|Company|Organization|Witness|Notary)\b",
        re.IGNORECASE,
    )
    body_text = re.compile(
        r"\b(?:agree|shall|each|party|parties|herein|pursuant|effective|"
        r"this\s+agreement|whereas|now\s+therefore)\b",
        re.IGNORECASE,
    )
    for m in re.finditer(r"\bBy\s*:|\bSignature\s*:|\bPrint\s+Name\s*:", text, re.IGNORECASE):
        window = text[m.end(): m.end() + 400]
        for line in window.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            if PLACEHOLDER_RE.search(stripped):
                continue
            if DOCUSIGN_FIELD_RE.search(stripped):
                continue
            if skip_labels.match(stripped):
                continue
            if re.match(r"^_{2,}", stripped):
                continue
            if len(stripped) > 80:
                continue
            if body_text.search(stripped):
                continue
            count += 1
            break
    return count


def detect_signatures(text: str, is_scanned: bool) -> tuple[str, str]:
    """
    Returns (has_evidence_str, signing_status_str).

    Signed   = strong confirmation: DocuSign timestamps, Adobe completed, or 2+ text blocks
    Review   = electronic platform detected but signing not confirmed from text
    Unsigned = readable doc, no signature evidence found
    """
    if is_scanned or not text.strip():
        return ("False", "Review")

    # ── DocuSign ──────────────────────────────────────────────────────────────
    if DOCUSIGN_DOC_RE.search(text):
        # Strategy 1: signer timestamps embedded at doc end
        if len(DOCUSIGN_TIMESTAMP_RE.findall(text)) >= 2:
            return ("True", "Signed")
        # Strategy 2: /s/ markers (some DocuSign configurations)
        if len(SLASH_S_RE.findall(text)) >= 2:
            return ("True", "Signed")
        # Strategy 3: text-based names in By: blocks (within 400-char window)
        if _count_filled_by_blocks(text) >= 2:
            return ("True", "Signed")
        # Strategy 4: dates clustered immediately after last Envelope ID
        # (some DocuSign configurations embed signer dates in locale formats, e.g. 20-Mar-2025)
        env_positions = [m.end() for m in DOCUSIGN_DOC_RE.finditer(text)]
        if env_positions:
            tail = text[env_positions[-1]: env_positions[-1] + 300]
            if len(DATE_ANY_RE.findall(tail)) >= 2:
                return ("True", "Signed")
        # DocuSign detected but signatures are graphic overlays — cannot verify from text
        return ("True", "Review")

    # ── Adobe Sign / PandaDoc / E-SIGNED platforms ────────────────────────────
    if ADOBE_AUDIT_RE.search(text):
        # Definitive completion signal
        if ADOBE_COMPLETED_RE.search(text):
            return ("True", "Signed")
        # 2+ e-signed events
        if len(ESIGNED_BY_RE.findall(text)) >= 2:
            return ("True", "Signed")
        # Certificate-style entries: "Name (Month D, YYYY HH:MM TZ)" — 2+ = Signed
        if len(ESIG_CERT_RE.findall(text)) >= 2:
            return ("True", "Signed")
        # Platform detected but completion not confirmed
        return ("True", "Review")

    # ── Non-electronic signatures ─────────────────────────────────────────────
    slash_count = len(SLASH_S_RE.findall(text))
    if slash_count >= 2:
        return ("True", "Signed")

    by_count = _count_filled_by_blocks(text)
    if by_count >= 2:
        return ("True", "Signed")

    evidence = slash_count >= 1 or by_count >= 1
    return ("True" if evidence else "False", "Unsigned")


def _normalize_date(raw: str) -> str | None:
    try:
        dt = dparser.parse(raw, dayfirst=False)
        # Sanity-check: reject dates before 1990 or after 2050
        if dt.year < 1990 or dt.year > 2050:
            return None
        return dt.strftime("%Y-%m-%d")
    except Exception:
        return None


def _trim_to_entity_name(raw: str) -> str:
    """Trim a party description (possibly with address/state) down to just the entity name."""
    raw = _clean_entity_name(raw)
    # Always try to strip state/type description suffixes, not just when >80 chars
    m = re.match(
        r"(.{5,}?)(?:,\s*a\s+(?:division|corporation|company|limited|partnership|"
        r"Delaware|Florida|Tennessee|California|Nevada|New York|Indiana|Ohio|Georgia|Texas|"
        r"Colorado|Illinois|Michigan|Virginia|Pennsylvania|Maryland|New Jersey|North Carolina)|"
        r",\s*an?\s+[\w\s]+?\s+(?:company|corp|corporation|partnership|liability|limited|ltd\.?)|"
        r",\s*located\s+at|,\s*of\s+\d|,\s*having\s+|with\s+its\s+principal\s+place|with\s+offices\s+at)",
        raw, re.IGNORECASE,
    )
    if m and len(m.group(1).strip()) >= 3:
        return m.group(1).strip(" ,.")
    if len(raw) <= 80:
        return raw
    # Fallback: stop at first comma
    first = raw.split(",")[0].strip()
    if len(first) >= 3:
        return first
    return raw[:80].strip()


_BY_AND_BETWEEN_RE = re.compile(r'^.*?by\s+and\s+between\s+', re.IGNORECASE)
_BETWEEN_RE        = re.compile(r'^.*?\bbetween\s+the\s+undersigned[,\s]+', re.IGNORECASE)
_PREAMBLE_RE       = re.compile(r'^(?:Page\s+\d+\s+of\s+\d+\s+)?(?:MUTUAL\s+)?(?:NON.?DISCLOSURE|CONFIDENTIALITY)\s+AGREEMENT\s+', re.IGNORECASE)

_GARBAGE_CANDIDATE_RE = re.compile(
    r'\bon\s+the\s+(one|other)\s+hand\b'
    r'|enter\s+into\s+this'
    r'|mutual\s+non.?disclosure'
    r'|non.?disclosure\s+agreement'
    r'|confidentiality\s+agreement'
    r'|this\s+agreement\b'
    r'|page\s+\d+\s+of'
    r'|between\s+the\s+undersigned'
    r'|^its\s+\w'               # "its Affiliates", "its parent", etc. — relative pronoun, not an entity
    r'|^the\s+(?:counterparty|party)\s+identified'  # template placeholder phrases
    r'|collectively.*parties'                        # "collectively, the 'Parties'" boilerplate
    r'|�',                  # Unicode replacement character — PDF encoding artifact
    re.IGNORECASE,
)


# DocuSign e-signature audit page ("Certificate Of Completion") — an artifact of
# the signing envelope, not the agreement itself.
_DOCUSIGN_CERT_RE = re.compile(r'Certificate\s+Of\s+Completion.{0,120}Envelope\s+Id', re.IGNORECASE)


def _clean_entity_name(raw: str) -> str:
    # Strip leading boilerplate before entity name
    raw = _BY_AND_BETWEEN_RE.sub('', raw)
    raw = _BETWEEN_RE.sub('', raw)
    raw = _PREAMBLE_RE.sub('', raw)
    # Strip leading punctuation/parens that aren't part of the name
    raw = re.sub(r'^[\s\("\'（,\.]+', '', raw)
    # Strip alias parens, collective/affiliate suffixes
    raw = re.sub(r'\s*[\("\'（][^)"\'）]{1,60}[\)"\'）]', "", raw)
    raw = re.sub(r"\s+and\s+its\s+(?:Affiliates|Subsidiaries|Related\s+Companies).*$", "", raw, flags=re.IGNORECASE)
    raw = re.sub(r"\s+its\s+(?:employees|officers|directors|representatives|affiliates|subsidiaries).*$", "", raw, flags=re.IGNORECASE)
    raw = re.sub(r"\s+together\s+with.*$", "", raw, flags=re.IGNORECASE)
    raw = re.sub(r"\s+\(collectively.*$", "", raw, flags=re.IGNORECASE)
    return raw.strip(" ,.()")


def extract_counterparty(text: str, vendor_folder: str = "") -> str | None:
    # Normalize line breaks for PDF text-flow artifacts before regex matching
    zone = re.sub(r"\s+", " ", text[:2000])

    # A DocuSign "Certificate Of Completion" page is an e-sign audit artifact, not
    # an agreement: it has no "by and between" clause, so every party-clause
    # strategy below misses and the entity-suffix scan (Strategy 4) fires on the
    # signer block instead — scooping the signer's JOB TITLE in with the company
    # ("Chief Operating Officer CONNORS AND ASSOCIATES, LLC"). Skip straight to
    # the vendor-folder fallback, which is the only reliable signal on these pages.
    if _DOCUSIGN_CERT_RE.search(zone):
        return vendor_folder.strip() if vendor_folder and len(vendor_folder.strip()) >= 2 else None

    def _valid_candidate(c: str) -> bool:
        """Return False if candidate is a garbage extraction or refers to DCS."""
        if not c or len(c) < 5:
            return False
        if re.match(r'^[_\s\-]+$', c):
            return False
        if 'DCS' in c or 'Designed Conveyor' in c or DCS_RE.search(c):  # catches all DCS variants
            return False
        if _GARBAGE_CANDIDATE_RE.search(c):
            return False
        return True

    # Strategy 0: DCS collective definition closure → "(collectively, "DCS")" then next ", and [PARTY]"
    # Handles: DCS...(collectively, "DCS") [address], and 6Sense Insights, Inc., together...
    m_dcs_close = re.search(r'"DCS"\)', zone, re.IGNORECASE)
    if m_dcs_close:
        after = zone[m_dcs_close.end():]
        m_party = re.search(
            r',\s+and\s+(.{3,120}?)'
            r'(?:,\s*(?:together|each|having|on\s+the|with\s+its)|\s*[\(\"\'\n\.])',
            after,
            re.IGNORECASE | re.DOTALL,
        )
        if m_party:
            candidate = _trim_to_entity_name(m_party.group(1))
            if _valid_candidate(candidate) and len(candidate) <= 80:
                return candidate

    # Strategy 1: DCS boilerplate "..., on the one hand, and [COUNTERPARTY]"
    m = DCS_ONE_HAND_RE.search(zone)
    if m:
        candidate = _trim_to_entity_name(m.group(1))
        if _valid_candidate(candidate) and len(candidate) <= 80:
            return candidate

    # Strategy 2: "and [ENTITY], on the other hand"
    m = OTHER_HAND_RE.search(zone)
    if m:
        candidate = _trim_to_entity_name(m.group(1))
        if _valid_candidate(candidate) and len(candidate) <= 80:
            return candidate

    # Strategy 3a: "between X, and Y" with comma separator (handles names containing "and")
    for m in PARTY_COMMA_AND_RE.finditer(zone):
        p1, p2 = m.group(1).strip(), m.group(2).strip()
        if DCS_RE.search(p1) and not DCS_RE.search(p2):
            candidate = _trim_to_entity_name(p2)
            if _valid_candidate(candidate) and len(candidate) <= 100:
                return candidate
        if DCS_RE.search(p2) and not DCS_RE.search(p1):
            candidate = _trim_to_entity_name(p1)
            if _valid_candidate(candidate) and len(candidate) <= 100:
                return candidate

    # Strategy 3b: generic "by and between X and Y"
    for m in PARTY_BLOCK_RE.finditer(zone):
        p1, p2 = m.group(1).strip(), m.group(2).strip()
        if DCS_RE.search(p1) and not DCS_RE.search(p2):
            candidate = _trim_to_entity_name(p2)
            if _valid_candidate(candidate) and len(candidate) <= 100:
                return candidate
        if DCS_RE.search(p2) and not DCS_RE.search(p1):
            candidate = _trim_to_entity_name(p1)
            if _valid_candidate(candidate) and len(candidate) <= 100:
                return candidate

    # Strategy 4: entity-suffix scan (non-DCS entity near top of doc)
    for em in ENTITY_SUFFIX_RE.finditer(zone):
        start = max(0, em.start() - 60)
        window = zone[start : em.end()]
        # Require actual uppercase start (no re.IGNORECASE on [A-Z] to avoid matching "by and between")
        nm = re.search(
            r"([A-Z][A-Za-z0-9\s,\.&\-']+?" + ENTITY_SUFFIX_RE.pattern + r")",
            window,
        )
        if nm:
            candidate = _trim_to_entity_name(nm.group(0))
            if _valid_candidate(candidate) and len(candidate) <= 80:
                return candidate

    # Strategy 5: VendorFolder fallback
    if vendor_folder and len(vendor_folder.strip()) >= 2:
        return vendor_folder.strip()

    return None


def extract_effective_date(text: str) -> str | None:
    zone = text[:2000]
    date_re = re.compile(DATE_PAT, re.IGNORECASE)

    # Primary: standard date immediately after trigger keyword
    for m in EFFECTIVE_DATE_RE.finditer(zone):
        raw = date_re.search(m.group(0))
        if raw:
            result = _normalize_date(raw.group(0))
            if result:
                return result

    # Fallback: "made as of the _30_ day of _April_____, 2020" template style
    for m in _AS_OF_THE_RE.finditer(zone):
        window = zone[m.end(): m.end() + 80]
        dm = _DAY_OF_RE.search(window)
        if dm:
            raw_date = f"{dm.group(1)} {dm.group(2).strip('_ ')} {dm.group(3)}"
            result = _normalize_date(raw_date)
            if result:
                return result
        # Also try a standard date in the window (e.g. "as of the April 30, 2020")
        raw = date_re.search(window)
        if raw:
            result = _normalize_date(raw.group(0))
            if result:
                return result

    return None


def extract_signature_dates(text: str) -> list[str]:
    """Find dates in signature blocks (latter half of doc) — EffectiveDate fallback."""
    date_re = re.compile(DATE_PAT, re.IGNORECASE)
    zone = text[len(text) // 2:]
    dates = []
    for m in SIG_DATE_RE.finditer(zone):
        raw = date_re.search(m.group(0))
        if raw:
            normalized = _normalize_date(raw.group(0))
            if normalized:
                dates.append(normalized)
    return sorted(set(dates))


def detect_at_will_termination(text: str) -> bool:
    return bool(AT_WILL_RE.search(text))


def extract_expiration_date(text: str, effective_iso: str | None) -> str | None:
    # Strategy A: explicit expiration date in text
    date_re = re.compile(DATE_PAT, re.IGNORECASE)
    for m in EXPIRATION_TRIGGER_RE.finditer(text):
        raw = date_re.search(m.group(0))
        if raw:
            result = _normalize_date(raw.group(0))
            if result:
                return result

    # Strategy B: N-year term arithmetic
    if effective_iso:
        for m in TERM_YEARS_RE.finditer(text):
            qty_str = m.group(1).lower()
            qty = WORD_TO_NUM.get(qty_str) or (int(qty_str) if qty_str.isdigit() else None)
            if qty and 1 <= qty <= 20:
                try:
                    eff = dparser.parse(effective_iso)
                    exp = eff + relativedelta(years=qty)
                    return exp.strftime("%Y-%m-%d")
                except Exception:
                    pass

    return None

# ── CSV helpers ───────────────────────────────────────────────────────────────

def load_csv(csv_path: Path) -> pd.DataFrame:
    return pd.read_csv(str(csv_path), dtype=str, keep_default_na=False)


ISO_DATE_RE = re.compile(r"\d{4}-\d{2}-\d{2}$")


def refresh_days_until_expiration(df: pd.DataFrame) -> None:
    """Recompute DaysUntilExpiration: whole days from today to ExpirationDate.

    Positive = days remaining, 0 = expires today, negative = days since it
    expired. Blank when ExpirationDate is not an ISO date (e.g. "on notice").
    Creates the column right after ExpirationDate if it does not exist yet.
    This is a derived value, so it is recomputed on every write.
    """
    today = date.today()

    def _days(raw: str) -> str:
        raw = (raw or "").strip()
        if not ISO_DATE_RE.match(raw):
            return ""
        try:
            return str((date.fromisoformat(raw) - today).days)
        except ValueError:
            return ""

    values = df["ExpirationDate"].apply(_days)
    if "DaysUntilExpiration" in df.columns:
        df["DaysUntilExpiration"] = values
    else:
        df.insert(df.columns.get_loc("ExpirationDate") + 1,
                  "DaysUntilExpiration", values)


def write_csv(df: pd.DataFrame, csv_path: Path, dry_run: bool = False):
    if dry_run:
        return
    refresh_days_until_expiration(df)
    tmp = csv_path.with_suffix(".csv.tmp")
    bak = csv_path.with_suffix(".csv.bak")
    df.to_csv(str(tmp), index=False, quoting=csv.QUOTE_ALL)
    shutil.copy2(str(csv_path), str(bak))
    tmp.replace(csv_path)


def _norm_path(p: str) -> str:
    return p.replace("\\", "/").strip()


def rows_for_filepath(df: pd.DataFrame, file_path_key: str) -> list[int]:
    key = _norm_path(file_path_key)
    mask = df["FilePath"].apply(_norm_path) == key
    return df.index[mask].tolist()


def resolve_abs_path(df: pd.DataFrame, row_idx: int, root: Path | None = None) -> Path | None:
    loc  = df.at[row_idx, "ContractLocation"]
    fp   = _norm_path(df.at[row_idx, "FilePath"])
    base = LOCATION_ROOTS.get(loc)
    if base is None:
        return None
    full = base / Path(PurePosixPath(fp))
    return full if full.exists() else None

# ── Target-row builders ───────────────────────────────────────────────────────

def build_target_indices(
    df: pd.DataFrame,
    positional: list[str],
    vendor: str | None,
    location: str | None,
    doc_type: str | None,
    all_rows: bool,
) -> list[int]:
    indices: set[int] = set()

    if positional:
        for key in positional:
            found = rows_for_filepath(df, key)
            if not found:
                print(f"  [WARN] No CSV row for FilePath: {key}")
            indices.update(found)

    if vendor:
        mask = df["VendorFolder"].str.contains(vendor, case=False, na=False)
        vendor_idx = set(df.index[mask].tolist())
        if not vendor_idx:
            print(f"  [WARN] No rows matched --vendor: \"{vendor}\"")
        if location:
            loc_mask = df["ContractLocation"] == location
            loc_idx  = set(df.index[loc_mask].tolist())
            indices.update(vendor_idx & loc_idx)
        else:
            indices.update(vendor_idx)
    elif location and not positional:
        mask = df["ContractLocation"] == location
        loc_idx = set(df.index[mask].tolist())
        if not loc_idx:
            print(f"  [WARN] No rows matched --location: \"{location}\"")
        indices.update(loc_idx)

    if doc_type and not positional and not vendor and not location:
        mask = df["DocType"].str.contains(doc_type, case=False, na=False)
        type_idx = set(df.index[mask].tolist())
        if not type_idx:
            print(f"  [WARN] No rows matched --type: \"{doc_type}\"")
        indices.update(type_idx)

    if all_rows and not positional and not vendor and not location and not doc_type:
        indices.update(df.index.tolist())

    return sorted(indices)

# ── Per-file scan ─────────────────────────────────────────────────────────────

def scan_file(
    df: pd.DataFrame,
    row_indices: list[int],
    abs_path: Path,
    vendor_folder: str,
    dry_run: bool,
    recheck_signing: bool = False,
    recheck_counterparty: bool = False,
) -> dict:
    """Scan one file, update df rows in-place. Returns result dict."""
    result = {
        "path": abs_path,
        "rows": row_indices,
        "changes": [],
        "status": "ok",
        "note": "",
    }

    text, is_scanned, skip_reason = extract_text(abs_path)

    if skip_reason:
        result["status"] = "skipped"
        result["note"]   = skip_reason
        return result

    # Detect DocType first — needed to suppress signing detection for non-contracts
    doc_type = detect_doc_type(text) if text else None

    # Word docs and non-contract DocTypes (Tax-Form, Exhibit, etc.) are always Unsigned
    if abs_path.suffix.lower() in (".docx", ".doc") or doc_type in _NON_CONTRACT_TYPES:
        has_kw, signing_status = "False", "Unsigned"
    else:
        # Signature detection must run before text is cleared for scanned PDFs
        has_kw, signing_status = detect_signatures(text, is_scanned)

    if is_scanned:
        result["status"] = "scanned_pdf"
        result["note"]   = "image-only PDF, no text extracted"
        text = ""  # clear so other extractors skip cleanly

    # Run remaining extractors (doc_type already detected above)
    counterparty = extract_counterparty(text, vendor_folder)
    eff_date     = extract_effective_date(text) if text else None
    if eff_date is None and text:
        sig_dates = extract_signature_dates(text)
        if sig_dates:
            eff_date = sig_dates[0]  # oldest date in signature blocks
    exp_date = extract_expiration_date(text, eff_date) if text else None
    if exp_date is None and text and detect_at_will_termination(text):
        exp_date = "upon written notice"

    extracted = {
        "DocType":          doc_type,
        "HasSignedKeyword": has_kw   if text else None,
        "SigningStatus":    signing_status if text else None,
        "CounterpartyName": counterparty,
        "EffectiveDate":    eff_date,
        "ExpirationDate":   exp_date,
    }

    # Write into df rows
    for idx in row_indices:
        for field, value in extracted.items():
            if value is None:
                continue
            old = df.at[idx, field]
            if old == value:
                continue

            # Signed and Unsigned are permanent — scanner never overwrites them.
            # Review is temporary and can be resolved to Signed or Unsigned.
            # --recheck-signing bypasses this for a one-time migration run.
            if field == 'SigningStatus' and not recheck_signing:
                if old in ('Signed', 'Unsigned') and old != value:
                    continue

            # CounterpartyName is fill-only. Extraction is heuristic and this is
            # the field humans most often hand-correct to the legal name on the
            # document; a rescan must not silently revert curated values (e.g. a
            # DocuSign certificate page falls back to the vendor-folder name).
            # Use --recheck-counterparty to force re-extraction.
            if field == 'CounterpartyName' and not recheck_counterparty:
                if old:
                    continue

            # DocType update rules: preserve specific types; allow narrowing upgrades.
            if field == 'DocType' and old:
                if old == 'Other':
                    pass  # Other is a catch-all; allow replacement
                elif old == 'NDA' and value == 'MNDA':
                    pass  # NDA→MNDA upgrade (mutual is more specific)
                elif old == 'License' and value == 'EULA':
                    pass  # License→EULA upgrade (end-user is more specific)
                else:
                    continue  # preserve all other existing DocTypes

            if not dry_run:
                df.at[idx, field] = value
            tag = "[DRY RUN] " if dry_run else ""
            result["changes"].append(
                f"    {tag}Row {idx} [{field}]: \"{old}\" -> \"{value}\""
            )

    return result

# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    # Ensure stdout can handle non-ASCII characters from PDF text (Windows cp1252 compat)
    if hasattr(sys.stdout, 'reconfigure'):
        sys.stdout.reconfigure(encoding='utf-8', errors='replace')

    parser = argparse.ArgumentParser(
        description="Scan contract files and update contract-catalog.csv",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument(
        "files",
        nargs="*",
        metavar="FILEPATH",
        help="One or more FilePath keys (e.g. \"Disney/EXHIBIT C.pdf\")",
    )
    parser.add_argument("--vendor",   metavar="NAME",   help="Scan all rows matching VendorFolder (partial, case-insensitive)")
    parser.add_argument("--location", metavar="FOLDER", help="Scan all rows in ContractLocation")
    parser.add_argument("--type",     metavar="TYPE",   help="Scan all rows matching DocType (partial, case-insensitive)")
    parser.add_argument("--all",      action="store_true", help="Scan entire catalog")
    parser.add_argument("--prune",    action="store_true", help="Remove catalog rows whose file no longer exists on disk")
    parser.add_argument("--dry-run",  action="store_true", help="Print changes without writing CSV")
    parser.add_argument("--recheck-signing", action="store_true",
                        help="Re-evaluate SigningStatus even for already-Signed/Unsigned rows (one-time migration)")
    parser.add_argument("--recheck-counterparty", action="store_true",
                        help="Re-extract CounterpartyName even for rows that already have one "
                             "(overwrites hand-curated legal names — one-time migration use only)")
    parser.add_argument("--csv",      metavar="PATH",   default=str(DEFAULT_CSV), help="Path to contract-catalog.csv")

    args = parser.parse_args()

    csv_path = Path(args.csv).resolve()
    if not csv_path.exists():
        sys.exit(f"CSV not found: {csv_path}")

    root = SHAREPOINT

    print("scan-contract.py  --  Contract Catalog Scanner")
    print("=" * 54)
    print(f"CSV:  {csv_path}")
    for _loc, _base in LOCATION_ROOTS.items():
        print(f"Root: [{_loc}] {_base}")
    if args.dry_run:
        print("Mode: DRY RUN (no changes written)")
    print()

    df = load_csv(csv_path)
    print(f"Loaded {len(df)} rows.\n")

    target_indices = build_target_indices(
        df,
        positional=args.files,
        vendor=args.vendor,
        location=args.location,
        doc_type=args.type,
        all_rows=args.all,
    )

    if not target_indices and not args.prune:
        sys.exit("No target rows resolved. Check your arguments.")

    # Deduplicate by FilePath so the same physical file isn't scanned twice
    seen_paths: set[str] = set()
    file_groups: list[tuple[str, list[int]]] = []
    for idx in target_indices:
        key = _norm_path(df.at[idx, "FilePath"])
        if key not in seen_paths:
            seen_paths.add(key)
            group_indices = rows_for_filepath(df, key)
            file_groups.append((key, group_indices))

    if file_groups:
        print(f"Scanning {len(file_groups)} file(s) across {len(target_indices)} row(s).\n")

    t_start = time.time()
    n_ok = n_skipped = n_scanned = n_missing = n_changed = 0

    for i, (fp_key, row_indices) in enumerate(file_groups, 1):
        vendor_folder = df.at[row_indices[0], "VendorFolder"] if row_indices else ""
        abs_path = resolve_abs_path(df, row_indices[0], root) if row_indices else None

        # Print progress every 10 files on large batches
        if len(file_groups) > 20 and i % 10 == 0:
            elapsed = time.time() - t_start
            print(f"  ... {i}/{len(file_groups)} files ({elapsed:.0f}s elapsed)")

        print(f"Scanning [{i}/{len(file_groups)}]: {fp_key}")

        if abs_path is None:
            print(f"  [NOT FOUND] File not on local disk (may be cloud-only / not synced via OneDrive) - skipped")
            n_missing += 1
            continue

        result = scan_file(df, row_indices, abs_path, vendor_folder, args.dry_run,
                           recheck_signing=args.recheck_signing,
                           recheck_counterparty=args.recheck_counterparty)

        if result["status"] == "skipped":
            print(f"  [SKIPPED] {result['note']}")
            n_skipped += 1
        elif result["status"] == "scanned_pdf":
            print(f"  [SCANNED PDF] {result['note']}")
            n_scanned += 1
        else:
            n_ok += 1

        if result["changes"]:
            n_changed += len(result["changes"])
            for line in result["changes"]:
                print(line)
        else:
            print("  (no changes)")

    # Prune orphaned rows (catalog entries with no matching file on disk)
    n_pruned = 0
    if args.prune:
        print("\n" + "=" * 54)
        print("Pruning orphaned catalog rows...")
        orphan_indices = []
        for idx in df.index:
            if resolve_abs_path(df, idx, root) is None:
                loc = df.at[idx, "ContractLocation"]
                fp  = df.at[idx, "FilePath"]
                print(f"  [ORPHAN] [{loc}] {fp}")
                orphan_indices.append(idx)
        n_pruned = len(orphan_indices)
        if n_pruned:
            if not args.dry_run:
                df = df.drop(index=orphan_indices).reset_index(drop=True)
                n_changed += n_pruned
            print(f"  {n_pruned} orphaned row(s) {'would be ' if args.dry_run else ''}removed.")
        else:
            print("  No orphaned rows found.")

    # Write CSV
    if not args.dry_run and n_changed > 0:
        write_csv(df, csv_path, dry_run=False)
        print(f"\nCSV written -> {csv_path}")
        print(f"Backup      -> {csv_path.with_suffix('.csv.bak')}")
    elif args.dry_run:
        print("\n[DRY RUN] CSV not written.")
    else:
        print("\nNo changes - CSV unchanged.")

    elapsed_total = time.time() - t_start
    print("\n" + "=" * 54)
    print(f"Summary: {len(file_groups)} file(s) scanned in {elapsed_total:.1f}s")
    print(f"  OK:          {n_ok}")
    print(f"  Skipped:     {n_skipped}")
    print(f"  Scanned PDF: {n_scanned}")
    print(f"  Not found:   {n_missing}")
    print(f"  Field changes: {n_changed}")
    print(f"  Pruned rows: {n_pruned}")
    if args.dry_run:
        print("  [DRY RUN - CSV not written]")


if __name__ == "__main__":
    main()
